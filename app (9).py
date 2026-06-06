# ============================================================
#   INSTASUMMARY
# ============================================================
import re, os, json, math, tempfile, logging, threading, time, datetime, random as _rnd
from collections import Counter

logging.basicConfig(level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s")
logger = logging.getLogger("InstaSummary")

import nltk, torch, gradio as gr, fitz, pandas as pd, numpy as np
from fpdf import FPDF, XPos, YPos
from langdetect import detect
from functools import lru_cache
from deep_translator import GoogleTranslator
from sklearn.feature_extraction.text import TfidfVectorizer
from transformers import (
    BartTokenizer, BartTokenizerFast, BartForConditionalGeneration,
    pipeline as hf_pipeline,
)

for _p in ["punkt","punkt_tab","stopwords","averaged_perceptron_tagger",
           "averaged_perceptron_tagger_eng","words"]:
    try: nltk.download(_p, quiet=True)
    except: pass
logger.info("Imports OK")

# ──────────────────────────────────────────────────────────────
#  STAR GENERATOR — full viewport 1920×1080
# ──────────────────────────────────────────────────────────────

def _star_css(n, mx=1920, my=1080, min_op=0.2, max_op=0.9):
    parts = []
    for _ in range(n):
        x    = _rnd.randint(1, mx)
        y    = _rnd.randint(1, my)
        blur = 0 if _rnd.random() < 0.7 else 1
        op   = round(_rnd.uniform(min_op, max_op), 2)
        parts.append(f"{x}px {y}px {blur}px rgba(255,255,255,{op})")
    return ",".join(parts)

STARS_A = _star_css(450, min_op=0.50, max_op=1.00)
STARS_B = _star_css(300, min_op=0.35, max_op=0.90)
STARS_C = _star_css(180, min_op=0.70, max_op=1.00)

# ──────────────────────────────────────────────────────────────
#  PRE-GENERATE WELCOME AUDIO
# ──────────────────────────────────────────────────────────────
_WELCOME_PATH = None

def _init_welcome():
    global _WELCOME_PATH
    try:
        from gtts import gTTS
        tts = gTTS(
            text="Welcome to InstaSummary — your intelligent AI-powered summarization assistant. "
                 "Paste text, upload a PDF, audio, or data file to get smart summaries instantly.",
            lang="en", slow=False
        )
        raw = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False).name
        tts.save(raw)
        try:
            from pydub import AudioSegment
            louder = AudioSegment.from_mp3(raw) + 10
            out    = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False).name
            louder.export(out, format="mp3", bitrate="128k")
            _WELCOME_PATH = out
        except:
            _WELCOME_PATH = raw
        logger.info(f"Welcome audio ready: {_WELCOME_PATH}")
    except Exception as e:
        logger.warning(f"Welcome audio failed: {e}")

threading.Thread(target=_init_welcome, daemon=True).start()

def _welcome_audio_fn():
    logger.info(f"WELCOME PATH = {_WELCOME_PATH}")
    return _WELCOME_PATH

# ──────────────────────────────────────────────────────────────
#  HISTORY
# ──────────────────────────────────────────────────────────────
HISTORY_FILE = "/tmp/instasummary_history.json"

def _load_hist():
    if not os.path.exists(HISTORY_FILE):
        return []
    try:
        with open(HISTORY_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []

def _save_hist(h):
    try:
        with open(HISTORY_FILE, "w", encoding="utf-8") as f:
            json.dump(h, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

# ──────────────────────────────────────────────────────────────
#  CONSTANTS
# ──────────────────────────────────────────────────────────────
LANGUAGES = {
    "English":"en","Hindi":"hi","Telugu":"te","Tamil":"ta",
    "Kannada":"kn","Malayalam":"ml","French":"fr","German":"de",
    "Spanish":"es","Japanese":"ja","Chinese":"zh-CN",
    "Arabic":"ar","Italian":"it","Korean":"ko","Portuguese":"pt",
}
GTTS_CODES = {
    "en":"en","hi":"hi","te":"te","ta":"ta","kn":"kn","ml":"ml",
    "fr":"fr","de":"de","es":"es","ja":"ja","zh-CN":"zh-TW",
    "ar":"ar","it":"it","ko":"ko","pt":"pt",
}
LANG_MAP = {
    "zh-cn":"zh-CN","zh-tw":"zh-CN","zh":"zh-CN","ja":"ja","ko":"ko",
    "ar":"ar","hi":"hi","te":"te","ta":"ta","kn":"kn","ml":"ml",
    "fr":"fr","de":"de","es":"es","pt":"pt","it":"it","ru":"ru","en":"en",
}
LENGTH_MAP = {
    "Short (50-80)":    (80,  50),
    "Medium (100-150)": (150, 100),
    "Long (200-280)":   (280, 200),
}
LANG_BONUS   = 40
LANG_CHOICES = list(LANGUAGES.keys())

# ──────────────────────────────────────────────────────────────
#  MODEL REGISTRY
# ──────────────────────────────────────────────────────────────
_MODELS: dict = {}
_WHISPER      = None

# ──────────────────────────────────────────────────────────────
#  BERT — sshleifer/distilbart-cnn-12-6
# ──────────────────────────────────────────────────────────────
def _get_bart_cnn():
    if "bert" not in _MODELS:
        name = "sshleifer/distilbart-cnn-12-6"
        logger.info("Loading DistilBART-CNN…")
        tok = BartTokenizerFast.from_pretrained(name)
        mdl = BartForConditionalGeneration.from_pretrained(
            name,
            torch_dtype=torch.float32,
            low_cpu_mem_usage=True,
        ).eval()
        mdl.to("cpu")
        logger.info("DistilBART ready")
        _MODELS["bert"] = (tok, mdl)
    return _MODELS["bert"]

# ──────────────────────────────────────────────────────────────
#  SEMANTIC RANKER — all-mpnet-base-v2
#  Integrated as Step 4 in the pipeline.
#  Ranks sentences by semantic importance before summarization.
# ──────────────────────────────────────────────────────────────
def _get_mpnet():
    """
    Lazy-loads all-mpnet-base-v2 and caches it.
    Returns a SentenceTransformer or a sentence-transformers pipeline.
    Falls back gracefully if the library is unavailable.
    """
    if "mpnet" not in _MODELS:
        try:
            from sentence_transformers import SentenceTransformer
            logger.info("Loading all-mpnet-base-v2 for semantic ranking…")
            model = SentenceTransformer("sentence-transformers/all-mpnet-base-v2")
            _MODELS["mpnet"] = model
            logger.info("all-mpnet-base-v2 ready")
        except Exception as e:
            logger.warning(f"sentence-transformers unavailable: {e}. "
                           "Falling back to TF-IDF ranking.")
            _MODELS["mpnet"] = None
    return _MODELS["mpnet"]


def _semantic_rank(sentences: list, top_k: int, doc_embedding=None) -> list:
    """
    Rank sentences by semantic relevance to the document centroid.
    Returns top_k sentences in their ORIGINAL ORDER (preserves flow).

    If all-mpnet-base-v2 is available: uses cosine similarity to doc centroid.
    Fallback: TF-IDF cosine similarity.
    """
    if not sentences:
        return sentences

    n = len(sentences)
    if n <= top_k:
        return sentences

    mpnet = _get_mpnet()

    if mpnet is not None:
        try:
            with torch.inference_mode():
                embs = mpnet.encode(sentences, convert_to_numpy=True,
                                    batch_size=32, show_progress_bar=False)
            # Document centroid = mean of all sentence embeddings
            centroid = embs.mean(axis=0)
            # Cosine similarity of each sentence to centroid
            norms    = np.linalg.norm(embs, axis=1, keepdims=True) + 1e-9
            c_norm   = np.linalg.norm(centroid) + 1e-9
            scores   = (embs / norms) @ (centroid / c_norm)

            # Pick top_k indices, then sort by original position
            top_idx = sorted(np.argsort(scores)[-top_k:].tolist())
            return [sentences[i] for i in top_idx]
        except Exception as e:
            logger.warning(f"Semantic ranking error: {e}. Falling back to TF-IDF.")

    # ── TF-IDF fallback ──────────────────────────────────────
    try:
        vec    = TfidfVectorizer(stop_words="english", max_features=512)
        mat    = vec.fit_transform(sentences).toarray()
        centroid = mat.mean(axis=0)
        norms  = np.linalg.norm(mat, axis=1, keepdims=True) + 1e-9
        c_norm = np.linalg.norm(centroid) + 1e-9
        scores = (mat / norms) @ (centroid / c_norm)
        top_idx = sorted(np.argsort(scores)[-top_k:].tolist())
        return [sentences[i] for i in top_idx]
    except Exception:
        # Last resort: return first top_k sentences
        return sentences[:top_k]


def _get_whisper():
    global _WHISPER
    if _WHISPER is None:
        logger.info("Loading Whisper…")
        _WHISPER = hf_pipeline("automatic-speech-recognition",
            model="openai/whisper-base", device=-1, chunk_length_s=30)
        logger.info("Whisper ready")
    return _WHISPER

def _preload():
    try:    _get_bart_cnn()
    except Exception as e: logger.error(f"DistilBART preload error: {e}")
    try:    _get_mpnet()
    except Exception as e: logger.warning(f"MPNet preload error: {e}")
    try:    _get_whisper()
    except Exception as e: logger.warning(f"Whisper preload error: {e}")

threading.Thread(target=_preload, daemon=True).start()

def _model_ready(_=None):
    return "bert" in _MODELS

# ──────────────────────────────────────────────────────────────
#  UTILITIES
# ──────────────────────────────────────────────────────────────
def _strip_emojis(text):
    if not text: return text
    return re.sub(r"[\U00010000-\U0010ffff]", "", text)

@lru_cache(maxsize=1024)
def _cached_detect(text):
    try:    return detect(text[:800])
    except: return "en"

def _dl(text):
    lang = _cached_detect(text)
    return LANG_MAP.get(lang.lower(), "en")

def _wc(t):
    return len(t.split())

def _sc(t):
    try:    return len(nltk.sent_tokenize(t))
    except: return max(1, len([s for s in t.split(".") if len(s.strip()) > 2]))

def _rt(n):
    wpm = 200; m = n / wpm
    if m < 1:   return f"{int(m * 60)} sec"
    if m < 60:  return f"{math.ceil(m)} min"
    return f"{int(m // 60)}h {math.ceil(m % 60)}min"

@lru_cache(maxsize=512)
def _cached_translate(chunk, tgt):
    try:    return GoogleTranslator(source="auto", target=tgt).translate(chunk)
    except: return chunk

def _translate(text, tgt, chunk_size=3500):
    if not text or not text.strip(): return text
    tgt = tgt.lower()
    if tgt in ("zh-cn","zh","chinese simplified"):       tgt = "chinese (simplified)"
    elif tgt in ("zh-tw","chinese traditional"):         tgt = "chinese (traditional)"
    chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)
              if text[i:i+chunk_size].strip()]
    return " ".join(_cached_translate(c, tgt) for c in chunks)

def _sp(f):
    if not f:            return None
    if isinstance(f, str): return f
    if isinstance(f, dict): return f.get("path") or f.get("name") or f.get("tmp_path")
    for a in ("path","name"):
        if hasattr(f, a): return getattr(f, a)
    return str(f)

def _clean(t):
    if not isinstance(t, str): t = str(t)
    t = t.encode("utf-8","ignore").decode("utf-8")
    t = _strip_emojis(t)
    return t

def _compress(original, summary):
    try:
        orig_wc = max(len(original.split()), 1)
        summ_wc = len(summary.split())
        return round((1 - summ_wc / orig_wc) * 100, 1)
    except: return 0.0

def _cn(code):
    for n, v in LANGUAGES.items():
        if v.lower() == code.lower(): return n
    return code.upper()

# ──────────────────────────────────────────────────────────────
#  STEP 2 — ENHANCED TEXT CLEANING
#  Removes noise while preserving multilingual content.
# ──────────────────────────────────────────────────────────────
def _preprocess(text: str) -> str:
    """
    Comprehensive text cleaning pipeline:
    - Removes page numbers, headers, footers, references
    - Removes OCR artifacts and broken characters
    - Deduplicates lines and removes near-duplicate sentences
    - Preserves paragraph structure and multilingual integrity
    """
    if not text or not text.strip():
        return text

    # ── Fix encoding artifacts ────────────────────────────────
    text = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]", " ", text)

    # ── Remove standalone page numbers ───────────────────────
    text = re.sub(r"(?m)^\s*[-–—]?\s*\d{1,4}\s*[-–—]?\s*$", "", text)

    # ── Remove citation/reference patterns ───────────────────
    text = re.sub(r"\[\d+(?:[,\s]\d+)*\]", "", text)           # [1], [2,3]
    text = re.sub(r"\(\s*[A-Za-z]+\s+\d{4}\s*\)", "", text)    # (Smith 2020)
    text = re.sub(r"\b(?:doi|DOI)\s*:\s*\S+", "", text)        # DOI links
    text = re.sub(r"https?://\S+", "", text)                    # URLs

    # ── Remove common header/footer patterns ─────────────────
    text = re.sub(
        r"(?im)^\s*(?:page|pg\.?|chapter|section|figure|table|fig\.?|"
        r"appendix|annex|contents|index|abstract|keywords?|"
        r"acknowledgem\w+|references?|bibliography)\s*[\d:.\-–]*\s*$",
        "", text
    )

    # ── Remove lines that look like running headers (ALL CAPS short lines) ──
    SQL_KW = {"SELECT","FROM","WHERE","JOIN","GROUP","ORDER","INSERT","UPDATE",
               "DELETE","CREATE","DROP","ALTER","HAVING","UNION","INDEX","INTO",
               "TABLE","SET","BY","ON","AND","OR","NOT","AS","IS","NULL","LEFT",
               "RIGHT","INNER","OUTER","FULL","CROSS","WITH","CASE","WHEN","THEN",
               "ELSE","END","DISTINCT","VALUES","PRIMARY","KEY","FOREIGN"}
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        # Skip very short ALL-CAPS lines (likely headers/footers)
        if (stripped and len(stripped) < 60 and stripped == stripped.upper()
                and len(stripped.split()) <= 6
                and not any(c.isdigit() for c in stripped[:3])):
            # But never strip SQL keywords
            words_upper = set(stripped.split())
            if not words_upper.issubset(SQL_KW):
                continue
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)

    # ── Remove excessive whitespace ───────────────────────────
    text = re.sub(r"[ \t]{3,}", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # ── Remove broken OCR artifacts ──────────────────────────
    # Lines with mostly non-alpha characters (e.g. "--- *** ---")
    lines = text.split("\n")
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue
        alpha_ratio = sum(1 for c in stripped if c.isalpha()) / max(len(stripped), 1)
        if alpha_ratio < 0.25 and len(stripped) < 80:
            continue  # skip artifact-heavy short lines
        if len(stripped) > 0 and len(stripped) <= 3:
            continue  # skip tiny fragments
        cleaned_lines.append(line)
    text = "\n".join(cleaned_lines)

    # ── Deduplicate consecutive identical lines ───────────────
    lines = text.split("\n")
    deduped, prev = [], None
    for line in lines:
        norm = line.strip().lower()
        if norm and norm == prev:
            continue
        deduped.append(line)
        prev = norm if norm else prev
    text = "\n".join(deduped)

    # ── Deduplicate near-identical sentences within paragraphs ──
    try:
        sents  = nltk.sent_tokenize(text)
        seen   = []
        unique = []
        for s in sents:
            norm = re.sub(r"\s+", " ", s.lower().strip())
            # Skip if very similar to a recent sentence (within last 5)
            is_dup = False
            for prev_s in seen[-5:]:
                # Simple Jaccard similarity on words
                w1 = set(norm.split())
                w2 = set(prev_s.split())
                if w1 and w2:
                    jaccard = len(w1 & w2) / len(w1 | w2)
                    if jaccard > 0.75:
                        is_dup = True
                        break
            if not is_dup:
                unique.append(s)
                seen.append(norm)
        text = " ".join(unique)
    except Exception:
        pass  # keep original text if dedup fails

    return text.strip()


# ──────────────────────────────────────────────────────────────
#  STEP 5 — INTELLIGENT CHUNKING
#  Paragraph-aware and sentence-aware chunking.
#  Avoids cutting sentences or concepts mid-way.
# ──────────────────────────────────────────────────────────────
def _split_chunks(text: str, max_w: int = 500) -> list:
    """
    Paragraph-aware, sentence-aware chunker.
    1. Tries to split on paragraph boundaries first.
    2. Falls back to sentence-level splitting within long paragraphs.
    3. Never cuts a sentence in half.
    """
    if not isinstance(text, str) or not text.strip():
        return [""]

    # ── Tokenize sentences ───────────────────────────────────
    try:
        all_sents = nltk.sent_tokenize(text)
    except Exception:
        all_sents = re.split(r"(?<=[.!?])\s+", text)

    # Filter out empty/tiny sentences
    all_sents = [s.strip() for s in all_sents if len(s.strip()) > 5]

    if not all_sents:
        return [text[:2000]]

    # ── Build paragraph-aware groups ─────────────────────────
    # We attempt to keep paragraph context together by grouping
    # sentences that appear in the same paragraph block.
    paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

    # If we have meaningful paragraphs, assign sentences to them
    sent_groups = []
    if len(paragraphs) > 1:
        for para in paragraphs:
            try:
                para_sents = nltk.sent_tokenize(para)
            except Exception:
                para_sents = re.split(r"(?<=[.!?])\s+", para)
            para_sents = [s.strip() for s in para_sents if len(s.strip()) > 5]
            if para_sents:
                sent_groups.append(para_sents)
    else:
        # No clear paragraphs — treat each sentence as its own group
        sent_groups = [[s] for s in all_sents]

    # ── Pack sentence groups into chunks ─────────────────────
    chunks, cur_sents, cur_wc = [], [], 0

    for group in sent_groups:
        group_wc = sum(len(s.split()) for s in group)

        # If adding this group exceeds limit and we have content, flush
        if cur_wc + group_wc > max_w and cur_sents:
            chunks.append(" ".join(cur_sents))
            cur_sents, cur_wc = [], 0

        # If a single paragraph group is itself > max_w, split it by sentence
        if group_wc > max_w:
            for sent in group:
                sw = len(sent.split())
                if cur_wc + sw > max_w and cur_sents:
                    chunks.append(" ".join(cur_sents))
                    cur_sents, cur_wc = [], 0
                cur_sents.append(sent)
                cur_wc += sw
        else:
            cur_sents.extend(group)
            cur_wc += group_wc

    if cur_sents:
        chunks.append(" ".join(cur_sents))

    # Final safety: filter empty chunks
    chunks = [c for c in chunks if c.strip()]
    return chunks if chunks else [text[:2000]]


# ──────────────────────────────────────────────────────────────
#  STEP 4 — SEMANTIC SENTENCE SELECTION
#  Uses all-mpnet-base-v2 to select the most important sentences
#  from the text before passing to DistilBART.
#  For large documents this dramatically improves quality.
# ──────────────────────────────────────────────────────────────
def _select_sentences_semantically(text: str, target_words: int) -> str:
    """
    Extract the most semantically important sentences from text,
    targeting approximately target_words words in the output.
    Preserves original sentence order for coherence.
    """
    try:
        sents = nltk.sent_tokenize(text)
    except Exception:
        sents = re.split(r"(?<=[.!?])\s+", text)

    sents = [s.strip() for s in sents if len(s.split()) >= 5]

    if not sents:
        return text

    total_words = sum(len(s.split()) for s in sents)

    # If already short enough, no selection needed
    if total_words <= target_words * 1.3:
        return text

    # Estimate how many sentences we need
    avg_words_per_sent = total_words / max(len(sents), 1)
    target_sents       = max(5, int(target_words / max(avg_words_per_sent, 1)))
    target_sents       = min(target_sents, len(sents))

    selected = _semantic_rank(sents, top_k=target_sents)
    return " ".join(selected)


# ──────────────────────────────────────────────────────────────
#  GENERATION — DistilBART, CPU-optimised
# ──────────────────────────────────────────────────────────────
def _run_bart(tok, mdl, text: str, mx: int, mn: int) -> str:
    """
    Run DistilBART inference with strong anti-repetition settings.
    Uses torch.inference_mode() for maximum CPU speed.
    """
    if not text or not text.strip():
        return ""

    inp = tok(
        text,
        return_tensors="pt",
        max_length=1024,
        truncation=True,
    )

    # Safety check: don't summarize if input is too short
    input_len = inp["input_ids"].shape[1]
    if input_len < 20:
        return tok.decode(inp["input_ids"][0], skip_special_tokens=True).strip()

    # Dynamically clamp max/min to avoid model errors
    safe_mx = min(mx, max(40, input_len - 2))
    safe_mn = min(mn, max(10, safe_mx // 2))

    with torch.inference_mode():
        ids = mdl.generate(
            inp["input_ids"],
            attention_mask=inp["attention_mask"],

            max_length=safe_mx,
            min_length=safe_mn,

            # num_beams=2: ~45% faster on CPU vs 4, negligible quality loss
            num_beams=2,
            no_repeat_ngram_size=3,
            repetition_penalty=2.2,
            length_penalty=1.0,

            early_stopping=True,
        )

    raw = tok.decode(ids[0], skip_special_tokens=True).strip()

    # ── Post-generation deduplication ────────────────────────
    return _dedup_summary(raw)


# ──────────────────────────────────────────────────────────────
#  STEP 7 — SUMMARY POST-PROCESSING (deduplication + cleanup)
# ──────────────────────────────────────────────────────────────
def _dedup_summary(text: str) -> str:
    """
    Removes repeated sentences and near-duplicate phrases from
    a generated summary. Ensures human-readable, clean output.
    """
    if not text:
        return text

    try:
        sents = nltk.sent_tokenize(text)
    except Exception:
        sents = re.split(r"(?<=[.!?])\s+", text)

    seen   = []
    unique = []
    for s in sents:
        s = s.strip()
        if not s:
            continue
        norm = re.sub(r"\s+", " ", s.lower())
        # Check Jaccard similarity against all previous unique sentences
        is_dup = False
        for prev in seen:
            w1 = set(norm.split())
            w2 = set(prev.split())
            if w1 and w2:
                jaccard = len(w1 & w2) / len(w1 | w2)
                if jaccard > 0.65:
                    is_dup = True
                    break
        if not is_dup:
            unique.append(s)
            seen.append(norm)

    result = " ".join(unique)

    # Remove runaway word repetitions like "technology technology technology"
    result = re.sub(r"\b(\w{4,})\s+(\1\s+){2,}", r"\1 ", result, flags=re.IGNORECASE)
    # Remove repeated short phrase pairs
    result = re.sub(r"(\b\w+\s+\w+\b)(\s+\1){2,}", r"\1", result, flags=re.IGNORECASE)

    return result.strip()


# ──────────────────────────────────────────────────────────────
#  STEP 6 — HIERARCHICAL SUMMARIZER (3-pass)
#  Pass 1: Summarise each semantic chunk
#  Pass 2: Merge & re-summarise if still long
#  Pass 3: Final refinement pass
# ──────────────────────────────────────────────────────────────
def _hierarchical(text: str, mx: int, mn: int) -> str:
    """
    Full 3-pass hierarchical summarization:
    - Pass 1: semantic ranking → chunk → summarise each chunk
    - Pass 2: merge chunk summaries → re-summarise if needed
    - Pass 3: final refinement → clean output

    Always uses DistilBART-CNN + all-mpnet-base-v2 semantic ranker.
    """
    if not text or not isinstance(text, str):
        return ""

    tok, mdl = _get_bart_cnn()
    wc       = len(text.split())

    # ── Step 4: Semantic sentence selection ──────────────────
    # For large docs, pre-select the most important sentences
    # using all-mpnet-base-v2 before chunking.
    # Target: ~3× the desired summary length for good coverage.
    semantic_target = mx * 3
    if wc > semantic_target * 1.5:
        logger.info(f"Semantic selection: {wc} → target ~{semantic_target} words")
        text = _select_sentences_semantically(text, target_words=semantic_target)
        wc   = len(text.split())
        logger.info(f"After semantic selection: {wc} words")

    # ── Step 5: Intelligent chunking ─────────────────────────
    chunks = _split_chunks(text, max_w=500) or [text]
    logger.info(f"Chunking: {len(chunks)} chunks | DistilBART")

    # ── Pass 1: Summarise each chunk ─────────────────────────
    pass1 = []
    for i, chunk in enumerate(chunks):
        chunk_wc = len(chunk.split())
        if chunk_wc < 20:
            pass1.append(chunk)
            continue
        # Scale chunk summary length proportionally
        chunk_mx = min(mx, max(50, chunk_wc // 3))
        chunk_mn = max(15, chunk_mx // 3)
        s = _run_bart(tok, mdl, chunk, chunk_mx, chunk_mn)
        if s.strip():
            pass1.append(s)
        logger.info(f"  Chunk {i+1}/{len(chunks)}: {chunk_wc}→{len(s.split())} words")

    combined = " ".join(p for p in pass1 if p.strip())

    # ── Pass 2: Re-summarise merged summaries if too long ─────
    if len(chunks) > 1 and len(combined.split()) > mx * 1.5:
        logger.info("Pass 2: re-summarising merged chunk summaries")
        # Semantic selection on merged summaries
        combined = _select_sentences_semantically(combined, target_words=mx * 2)
        pass2_chunks = _split_chunks(combined, max_w=400) or [combined]
        pass2 = []
        for chunk in pass2_chunks:
            chunk_wc = len(chunk.split())
            if chunk_wc < 20:
                pass2.append(chunk)
                continue
            chunk_mx = min(mx, max(50, chunk_wc // 2))
            chunk_mn = max(15, chunk_mx // 3)
            s = _run_bart(tok, mdl, chunk, chunk_mx, chunk_mn)
            if s.strip():
                pass2.append(s)
        combined = " ".join(p for p in pass2 if p.strip())

    # ── Pass 3: Final refinement pass ────────────────────────
    final_wc = len(combined.split())
    if final_wc > mx:
        logger.info(f"Pass 3: final refinement ({final_wc} → {mx} words)")
        # One more semantic selection to get the best sentences
        combined = _select_sentences_semantically(combined, target_words=mx)
        ref_mx   = mx
        ref_mn   = mn
        refined  = _run_bart(tok, mdl, combined, ref_mx, ref_mn)
        if refined.strip():
            combined = refined

    # ── Final deduplication ───────────────────────────────────
    combined = _dedup_summary(combined)

    return combined


# ──────────────────────────────────────────────────────────────
#  STEP 11 — SUMMARY LENGTH CONTROL
#  Ensures Short / Medium / Long produce NOTICEABLY different output.
# ──────────────────────────────────────────────────────────────
def _fmt_summary(text: str, fmt: str, length: str) -> str:
    """
    Enforces length targets post-generation.
    Short:  hard-trims to ~80 words
    Medium: ~150 words
    Long:   up to 280 words
    """
    text = text.strip()
    if not text:
        return text

    try:
        sents = nltk.sent_tokenize(text)
    except Exception:
        sents = re.split(r"(?<=[.!?])\s+", text)

    sents = [s.strip() for s in sents if s.strip()]
    if not sents:
        return text

    # Bullet format
    if "bullet" in fmt.lower():
        limit = {"Short (50-80)": 4, "Medium (100-150)": 6, "Long (200-280)": 9}.get(length, 6)
        return "\n".join(f"• {s}" for s in sents[:limit])

    # Paragraph format — enforce word limits strictly
    word_limits = {
        "Short (50-80)":    80,
        "Medium (100-150)": 150,
        "Long (200-280)":   280,
    }
    word_min = {
        "Short (50-80)":    40,
        "Medium (100-150)": 80,
        "Long (200-280)":   160,
    }

    target_max = word_limits.get(length, 150)
    target_min = word_min.get(length, 80)

    # Build result respecting the word limit
    result_sents, word_count = [], 0
    for s in sents:
        sw = len(s.split())
        if word_count + sw > target_max and word_count >= target_min:
            break
        result_sents.append(s)
        word_count += sw

    # If we have too few words, include more sentences
    if word_count < target_min and len(result_sents) < len(sents):
        for s in sents[len(result_sents):]:
            result_sents.append(s)
            word_count += len(s.split())
            if word_count >= target_min:
                break

    return " ".join(result_sents) if result_sents else text


def _apply_tone(text, tone):
    if tone != "Simple English": return text.strip()
    word_map = {
        "utilize":"use","approximately":"about","demonstrate":"show",
        "consequently":"so","however":"but","therefore":"so",
        "additionally":"also","facilitate":"help","numerous":"many",
        "obtain":"get","commence":"start","terminate":"end",
        "assist":"help","require":"need","purchase":"buy",
        "improve":"make better","commonly":"usually","frequently":"often",
        "rapidly":"quickly","initial":"first","final":"last",
        "attempt":"try","complete":"finish","indicate":"show",
        "explain":"tell","important":"key","significant":"important",
        "reduce":"decrease","increase":"grow","calculate":"find",
        "identify":"find","modify":"change","optimize":"improve",
        "select":"choose","eliminate":"remove","create":"make",
        "construct":"build","investigate":"look into","observe":"see",
        "achieve":"reach","maintain":"keep","compare":"compare",
        "analyze":"study","conclude":"end","prioritize":"focus on",
        "distribute":"give out","implement":"use","generate":"make",
        "develop":"build","establish":"set up","support":"help",
        "enhance":"improve","clarify":"explain","predict":"guess",
        "respond":"reply","transform":"change",
    }
    def replace(match):
        word = match.group(0)
        return word_map.get(word.lower(), word)
    pattern = r"\b(" + "|".join(map(re.escape, word_map.keys())) + r")\b"
    return re.sub(pattern, replace, text, flags=re.IGNORECASE).strip()


# ──────────────────────────────────────────────────────────────
#  ENHANCED CODE ANALYZER
#  Supports 14 languages with intelligent natural-language output.
#  Detects algorithms, patterns, frameworks, and code purpose.
# ──────────────────────────────────────────────────────────────

def _detect_code(text: str) -> str | None:
    """
    Detects programming language from source code.
    Supports: Python, Java, JavaScript, C, C++, C#, SQL,
              HTML, CSS, PHP, Go, Rust, Kotlin, Swift.
    Returns language name string or None if not code.
    """
    # ── Prose guard: if mostly natural language, skip ────────
    PROSE = {"the","is","are","was","were","this","that","have","has","will",
             "with","from","they","their","which","what","when","how","but",
             "and","for","not","you","your","can","may","in","of","to","a",
             "an","at","by","on","as","its","been","be","do","did","had"}
    t_low   = text.lower()
    words   = set(re.findall(r"[a-z]{3,}", t_low[:3000]))
    code_lines = [l for l in text.split("\n")
                  if re.search(r"[(){};:=<>\[\]]", l) and l.strip()]
    prose_score = len(words & PROSE)
    # High prose + low code lines = natural text
    if prose_score >= 14 and len(code_lines) < 4:
        return None

    sc = {}

    # ── Python ───────────────────────────────────────────────
    py = 0
    if re.search(r"\bdef\s+\w+\s*\(", text):          py += 4
    if re.search(r"\bclass\s+\w+[:(]", text):         py += 3
    if re.search(r"^from\s+\w|^import\s+\w", text, re.M): py += 3
    if "if __name__" in t_low:                         py += 4
    if re.search(r":\s*$", text, re.M):               py += 2
    if re.search(r"^\s{4,}\S", text, re.M):           py += 2
    if any(k in t_low for k in ["numpy","pandas","matplotlib","sklearn","torch",
                                  "tensorflow","keras","flask","django","fastapi"]): py += 3
    sc["Python"] = py

    # ── Java ─────────────────────────────────────────────────
    jv = 0
    if "public class" in t_low:                        jv += 5
    if "public static void main" in t_low:             jv += 5
    if re.search(r"\bSystem\.out\.print", text):       jv += 4
    if re.search(r"\bprivate|protected|public\b", t_low): jv += 2
    if re.search(r"\bnew\s+\w+\s*\(", text):          jv += 2
    if re.search(r"@Override|@Autowired|@SpringBoot", text): jv += 3
    if re.search(r"import\s+java\.|import\s+org\.", text): jv += 4
    sc["Java"] = jv

    # ── JavaScript ───────────────────────────────────────────
    js = 0
    if re.search(r"\bconst\s+\w+\s*=", text):         js += 3
    if re.search(r"\blet\s+\w+\s*=", text):           js += 2
    if "console.log" in t_low:                         js += 3
    if re.search(r"=>\s*[{(]", text):                 js += 3
    if re.search(r"\bfunction\s+\w+\s*\(", text):     js += 3
    if re.search(r"document\.|window\.|getElementById", text): js += 4
    if re.search(r"require\s*\(|module\.exports", text): js += 3
    if re.search(r"async\s+function|await\s+", text): js += 2
    if any(k in t_low for k in ["react","vue","angular","jquery","axios","express"]): js += 3
    sc["JavaScript"] = js

    # ── C ────────────────────────────────────────────────────
    c = 0
    if "#include" in t_low:                            c += 3
    if re.search(r"\bprintf\s*\(", text):             c += 4
    if re.search(r"\bscanf\s*\(", text):              c += 3
    if re.search(r"\bint\s+main\s*\(", text):         c += 5
    if re.search(r"\bmalloc\s*\(|\bfree\s*\(", text): c += 4
    if re.search(r"\bstruct\s+\w+", text):            c += 2
    # Differentiate C from C++: pure C has no cout or class
    if "cout" not in t_low and "class" not in t_low:  c += 2
    sc["C"] = c

    # ── C++ ──────────────────────────────────────────────────
    cpp = 0
    if "#include" in t_low:                            cpp += 2
    if re.search(r"\bcout\s*<<", text):               cpp += 5
    if re.search(r"\bcin\s*>>", text):                cpp += 4
    if re.search(r"\bstd::", text):                   cpp += 4
    if re.search(r"\bclass\s+\w+", text):             cpp += 3
    if re.search(r"\btemplate\s*<", text):            cpp += 4
    if re.search(r"\bvector<|map<|unordered_map<", text): cpp += 3
    if re.search(r"::|\bnamespace\b", text):          cpp += 2
    sc["C++"] = cpp

    # ── SQL ──────────────────────────────────────────────────
    sql = 0
    if re.search(r"\bSELECT\b.{1,200}\bFROM\b", text, re.I|re.S): sql += 6
    if re.search(r"\bCREATE\s+TABLE\b", text, re.I): sql += 5
    if re.search(r"\bINSERT\s+INTO\b", text, re.I):  sql += 5
    if re.search(r"\bUPDATE\b.{1,100}\bSET\b", text, re.I): sql += 5
    if re.search(r"\bDELETE\s+FROM\b", text, re.I):  sql += 4
    if re.search(r"\bWHERE\b|\bGROUP\s+BY\b|\bORDER\s+BY\b|\bJOIN\b", text, re.I): sql += 3
    if re.search(r"\bHAVING\b|\bUNION\b|\bINDEX\b", text, re.I): sql += 2
    sc["SQL"] = sql

    # ── HTML ─────────────────────────────────────────────────
    html = 0
    if re.search(r"<!DOCTYPE\s+html", text, re.I):   html += 6
    if re.search(r"<html[\s>]", text, re.I):         html += 5
    if re.search(r"<head[\s>]|<body[\s>]", text, re.I): html += 4
    if re.search(r"<div|<p>|<h[1-6]|<span|<a\s", text, re.I): html += 3
    if re.search(r"<form|<input|<button|<table", text, re.I): html += 2
    sc["HTML"] = html

    # ── Pick the best match — only 7 supported languages ─────
    if not sc: return None
    best  = max(sc, key=sc.get)
    score = sc[best]

    # Minimum threshold: need at least 4 points to confidently call it code
    if score < 4:
        return None

    # Resolve C vs C++ ambiguity: C++ wins if cout/std present
    if best == "C" and sc.get("C++", 0) >= sc["C"]:
        best = "C++"

    return best


# ──────────────────────────────────────────────────────────────
#  INTELLIGENT CODE EXPLANATION ENGINE
# ──────────────────────────────────────────────────────────────

def _analyze_code_structure(text: str, lang: str) -> dict:
    """
    Extracts detailed structural metrics from source code.
    Returns a dict with counts and detected patterns.
    """
    lines    = text.strip().split("\n")
    nonempty = [l for l in lines if l.strip()]
    t_low    = text.lower()
    result   = {
        "lang": lang, "total_lines": len(nonempty),
        "functions": [], "classes": [], "imports": [],
        "loops": 0, "branches": 0,
        "patterns": [], "frameworks": [], "purpose_hints": []
    }

    # ── Line counts ───────────────────────────────────────────
    cx = ("large module"  if len(nonempty) > 100 else
          "medium script" if len(nonempty) > 40  else
          "short script"  if len(nonempty) > 15  else
          "small utility")
    result["size_label"] = cx

    # ── Function extraction per language ─────────────────────
    func_patterns = {
        "Python":     r"^\s*def\s+([A-Za-z_]\w*)\s*\(",
        "Java":       r"(?:public|private|protected|static)[\w\s<>\[\]]+\s+([A-Za-z_]\w*)\s*\([^)]*\)\s*(?:throws\s+\w+\s*)?\{",
        "JavaScript": r"(?:function\s+([A-Za-z_]\w*)\s*\(|(?:const|let|var)\s+([A-Za-z_]\w*)\s*=\s*(?:async\s*)?\(?.*\)?\s*=>)",
        "C":          r"^\w[\w\s\*]+\s+([A-Za-z_]\w*)\s*\([^)]*\)\s*\{",
        "C++":        r"^\w[\w\s\*:~]+\s+([A-Za-z_]\w*)\s*\([^)]*\)\s*(?:const\s*)?\{",
        "SQL":        None,
        "HTML":       None,
    }
    fp = func_patterns.get(lang)
    if fp:
        for l in lines:
            m = re.match(fp, l)
            if m:
                name = next((g for g in m.groups() if g), None)
                if name and name not in ("if","for","while","switch","catch","main"):
                    result["functions"].append(name)
    result["functions"] = list(dict.fromkeys(result["functions"]))[:10]

    # ── Class extraction ──────────────────────────────────────
    class_patterns = {
        "Python":     r"^\s*class\s+([A-Za-z_]\w*)\s*[:(]",
        "Java":       r"(?:public|private|abstract|final)?\s*class\s+([A-Za-z_]\w*)",
        "JavaScript": r"^\s*class\s+([A-Za-z_]\w*)",
        "C++":        r"^\s*class\s+([A-Za-z_]\w*)",
        "C":          None,
        "SQL":        None,
        "HTML":       None,
    }
    cp = class_patterns.get(lang)
    if cp:
        for l in lines:
            m = re.search(cp, l)
            if m: result["classes"].append(m.group(1))
    result["classes"] = list(dict.fromkeys(result["classes"]))[:8]

    # ── Import extraction ─────────────────────────────────────
    import_patterns = {
        "Python":     r"^(?:import|from)\s+([\w.]+)",
        "Java":       r"^import\s+([\w.]+)",
        "JavaScript": r"(?:require\s*\(['\"]|from\s+['\"])([^'\"]+)",
        "C":          r"#include\s*[<\"]([\w./]+)[>\"]",
        "C++":        r"#include\s*[<\"]([\w./]+)[>\"]",
        "SQL":        None,
        "HTML":       None,
    }
    ip = import_patterns.get(lang)
    if ip:
        for l in lines:
            m = re.search(ip, l.strip())
            if m:
                raw = m.group(1)
                # Clean up: take root module only
                pkg = raw.split(".")[0].split("/")[-1].split(":")[0]
                if pkg and len(pkg) > 1: result["imports"].append(pkg)
    result["imports"] = list(dict.fromkeys(result["imports"]))[:10]

    # ── Loop and branch counting ──────────────────────────────
    loop_kw = {
        "Python":     r"\b(for|while)\b",
        "Java":       r"\b(for|while|do)\b",
        "JavaScript": r"\b(for|while|forEach|map|filter|reduce)\b",
        "C":          r"\b(for|while|do)\b",
        "C++":        r"\b(for|while|do)\b",
        "SQL":        "",
        "HTML":       "",
    }.get(lang, r"\b(for|while)\b")
    branch_kw = r"\b(if|elif|else|switch|case|guard|when|catch|try|except|unless)\b"
    if loop_kw:
        result["loops"]    = sum(1 for l in lines if re.search(loop_kw, l))
    result["branches"] = sum(1 for l in lines if re.search(branch_kw, l))

    # ── Pattern detection ─────────────────────────────────────
    patterns = []

    # Recursion
    if result["functions"]:
        for fn in result["functions"]:
            if text.count(fn + "(") > 1:
                patterns.append("recursion")
                break

    # Sorting
    if re.search(r"\bsort\b|\bsorted\b|\bquicksort\b|\bmergesort\b|\bbubble.?sort\b", t_low):
        patterns.append("sorting algorithm")

    # Searching
    if re.search(r"\bbinary.?search\b|\blinear.?search\b|\bsearch\b", t_low):
        if "search" in t_low and any(f in t_low for f in result["functions"] + ["find","lookup","search"]):
            patterns.append("search algorithm")

    # OOP
    if result["classes"]:
        patterns.append("object-oriented programming")

    # File I/O
    if re.search(r"\bopen\s*\(|\bfread\b|\bfwrite\b|FileReader|readFile|writeFile|\.read\(|\.write\(|fopen|fclose|BufferedReader", text):
        patterns.append("file handling")

    # Database / SQL queries
    if re.search(r"SELECT|INSERT|UPDATE|DELETE|cursor\.|execute\(|query\(|\.sql", text, re.I):
        patterns.append("database operations")

    # Web / HTTP
    if re.search(r"fetch\(|axios|XMLHttpRequest|requests\.get|requests\.post|http\.|HttpClient|@app\.route|@Get|@Post", text):
        patterns.append("HTTP/API calls")

    # ML/AI
    if re.search(r"fit\s*\(|predict\s*\(|train_test_split|model\.|\.compile\(|\.backward\(|gradient|loss\s*=|optimizer", text):
        patterns.append("machine learning")

    # Exception handling
    if re.search(r"\btry\b|\bcatch\b|\bexcept\b|\bfinally\b", t_low):
        patterns.append("exception handling")

    # Concurrency
    if re.search(r"\bthread\b|\basync\b|\bawait\b|\bgoroutine\b|\bcoroutine\b|\bcompletablefuture\b", t_low):
        patterns.append("asynchronous/concurrent execution")

    # Functional programming
    if re.search(r"\bmap\s*\(|\bfilter\s*\(|\breduce\s*\(|\blambda\b", t_low):
        patterns.append("functional programming patterns")

    # Data structures
    if re.search(r"\bqueue\b|\bstack\b|\blinkedlist\b|\bdeque\b|\btree\b|\bgraph\b|\bheap\b", t_low):
        patterns.append("data structure implementation")

    result["patterns"] = list(dict.fromkeys(patterns))

    # ── Framework detection ───────────────────────────────────
    fw_map = {
        "Python": {
            "flask":"Flask web framework","django":"Django web framework",
            "fastapi":"FastAPI framework","numpy":"NumPy","pandas":"Pandas",
            "matplotlib":"Matplotlib","sklearn":"scikit-learn",
            "torch":"PyTorch","tensorflow":"TensorFlow","keras":"Keras",
            "requests":"requests library","asyncio":"asyncio",
        },
        "Java": {
            "spring":"Spring Framework","hibernate":"Hibernate ORM",
            "junit":"JUnit testing","jackson":"Jackson JSON","lombok":"Lombok",
        },
        "JavaScript": {
            "react":"React","vue":"Vue.js","angular":"Angular",
            "express":"Express.js","jquery":"jQuery","axios":"Axios",
        },
        "C++": {
            "boost":"Boost library","qt":"Qt framework",
            "opencv":"OpenCV","eigen":"Eigen",
        },
        "C":    {},
        "SQL":  {},
        "HTML": {},
    }
    fw_lang = fw_map.get(lang, {})
    for kw, label in fw_lang.items():
        if kw in t_low:
            result["frameworks"].append(label)
    result["frameworks"] = result["frameworks"][:5]

    return result


def _generate_code_narrative(info: dict, text: str) -> str:
    """
    Generates a professional natural-language explanation of the code.
    Avoids generic statements; uses actual code content for specificity.
    """
    lang     = info["lang"]
    t_low    = text.lower()
    funcs    = info["functions"]
    classes  = info["classes"]
    imports  = info["imports"]
    loops    = info["loops"]
    branches = info["branches"]
    patterns = info["patterns"]
    fws      = info["frameworks"]
    lines    = info["total_lines"]
    size     = info["size_label"]

    parts = []

    # ── SQL narrative ─────────────────────────────────────────
    if lang == "SQL":
        ops = []
        if re.search(r"\bSELECT\b", text, re.I): ops.append("queries")
        if re.search(r"\bINSERT\b", text, re.I): ops.append("inserts records")
        if re.search(r"\bUPDATE\b", text, re.I): ops.append("updates data")
        if re.search(r"\bDELETE\b", text, re.I): ops.append("deletes records")
        if re.search(r"\bCREATE\s+TABLE\b", text, re.I): ops.append("creates tables")

        tables = re.findall(r"\bFROM\s+(\w+)|\bJOIN\s+(\w+)|\bINTO\s+(\w+)|\bUPDATE\s+(\w+)", text, re.I)
        table_names = list(dict.fromkeys([t for grp in tables for t in grp if t]))[:5]

        joins = re.findall(r"\b(INNER|LEFT|RIGHT|FULL|CROSS)\s+JOIN\b", text, re.I)
        has_agg = bool(re.search(r"\b(COUNT|SUM|AVG|MAX|MIN)\s*\(", text, re.I))
        has_group = bool(re.search(r"\bGROUP\s+BY\b", text, re.I))
        has_order = bool(re.search(r"\bORDER\s+BY\b", text, re.I))
        has_where = bool(re.search(r"\bWHERE\b", text, re.I))

        if ops:
            parts.append(f"This SQL script {', '.join(ops[:3])}.")
        if table_names:
            parts.append(f"It operates on the following table(s): {', '.join(table_names)}.")
        if joins:
            parts.append(f"It performs {len(joins)} JOIN operation(s) to combine related data.")
        if has_where:
            parts.append("Filtering conditions are applied using a WHERE clause to narrow the result set.")
        if has_agg:
            parts.append("Aggregate functions are used to compute summary statistics across rows.")
        if has_group:
            parts.append("Results are grouped to enable aggregation by category.")
        if has_order:
            parts.append("The output is sorted using an ORDER BY clause.")
        return " ".join(parts)

    # ── HTML narrative ────────────────────────────────────────
    if lang == "HTML":
        tags = re.findall(r"<(\w+)[\s>]", text, re.I)
        tag_counts = Counter(t.lower() for t in tags)
        top_tags = [t for t, _ in tag_counts.most_common(6) if t not in ("html","head","body","meta","link","script","style")]

        has_form = bool(re.search(r"<form", text, re.I))
        has_table = bool(re.search(r"<table", text, re.I))
        has_nav = bool(re.search(r"<nav|<header|<footer", text, re.I))
        has_js = bool(re.search(r"<script", text, re.I))
        has_css_link = bool(re.search(r"<link.+stylesheet|<style", text, re.I))

        parts.append(f"This HTML document defines the structure of a web page.")
        if has_nav:
            parts.append("It includes navigational elements such as headers or footers for page layout.")
        if top_tags:
            parts.append(f"The document uses {', '.join(top_tags[:4])} elements to organise content.")
        if has_form:
            parts.append("A form is present for collecting user input.")
        if has_table:
            parts.append("Tabular data is displayed using an HTML table.")
        if has_css_link:
            parts.append("External or embedded CSS is applied for visual styling.")
        if has_js:
            parts.append("JavaScript is embedded to add interactive behaviour.")
        return " ".join(parts)

    # ── Machine Learning narrative (Python-only) ─────────────
    if "machine learning" in patterns:
        ds = []
        if re.search(r"read_csv|load_data|dataset|pd\.read|ImageDataGenerator", t_low): ds.append("loads a dataset")
        if re.search(r"train_test_split|\.split\(", t_low): ds.append("splits data into training and test sets")

        model_type = "machine learning model"
        if re.search(r"LinearRegression|Ridge|Lasso", text): model_type = "linear regression model"
        elif re.search(r"RandomForest|ExtraTrees|GradientBoosting", text): model_type = "ensemble/forest model"
        elif re.search(r"SVM|SVC|SVR", text): model_type = "support vector machine"
        elif re.search(r"KNeighbors|KNN", text): model_type = "K-nearest neighbours classifier"
        elif re.search(r"Sequential|Dense|Conv2D|LSTM|GRU|Transformer", text): model_type = "neural network"
        elif re.search(r"XGBoost|LGBMClassifier|CatBoost", text): model_type = "gradient boosting model"
        elif re.search(r"KMeans|DBSCAN|AgglomerativeClustering", text): model_type = "clustering model"
        elif re.search(r"LogisticRegression", text): model_type = "logistic regression classifier"

        has_fit      = bool(re.search(r"\.fit\s*\(", text))
        has_predict  = bool(re.search(r"\.predict\s*\(", text))
        has_evaluate = bool(re.search(r"accuracy_score|classification_report|confusion_matrix|mean_squared_error|r2_score", t_low))

        if ds:
            parts.append(f"This script {' and '.join(ds)}.")
        parts.append(f"It builds and trains a {model_type}.")
        if has_fit:
            parts.append("The model is trained using a fit() call on the training data.")
        if has_predict:
            parts.append("Predictions are generated on new or test data using predict().")
        if has_evaluate:
            parts.append("Model performance is evaluated using standard metrics.")
        if fws:
            parts.append(f"Libraries used include {', '.join(fws[:3])}.")
        return " ".join(parts)

    
    # ── Smart purpose detection ─────────────────────────────

    if lang == "Python":

        if "fibonacci" in t_low:
            parts.append(
            "This Python program generates Fibonacci numbers using recursion and prints the sequence."
            )

        elif re.search(r"\bsort|sorted\b", t_low):
            parts.append(
            "This Python program sorts data and displays the ordered results."
            )

        elif re.search(r"\bsearch\b", t_low):
            parts.append(
            "This Python program searches for values within a dataset."
            )

        elif re.search(r"\bopen\(|read\(|write\(", text):
            parts.append(
                "This Python program reads from or writes to files."
        )
    
        else:
            parts.append(
                f"This Python program contains {len(funcs)} function(s), "
                f"{loops} loop(s), and {branches} conditional branch(es)."
            )

    elif lang == "Java":

        if "system.out.print" in t_low:
            parts.append(
                "This Java program prints output to the console and executes program logic."
            )

        elif re.search(r"\bsort\b", t_low):
            parts.append(
                "This Java program sorts data and displays the results."
            )

        else:
            parts.append(
                f"This Java program contains {len(funcs)} method(s), "
                f"{loops} loop(s), and {branches} conditional branch(es)."
            )
    
    elif lang == "C":

        parts.append(
            f"This C program contains {len(funcs)} function(s), "
            f"{loops} loop(s), and {branches} conditional branch(es)."
        )

    elif lang == "C++":

        parts.append(
            f"This C++ program contains {len(funcs)} function(s), "
            f"{loops} loop(s), and {branches} conditional branch(es)."
        )

    elif lang == "SQL":
    
        parts.append(
            "This SQL query retrieves, filters, or modifies database records."
        )

    elif lang == "HTML":

        parts.append(
            "This HTML document defines the structure and content of a web page."
        )

    else:

        parts.append(
            f"This {lang} program contains {len(funcs)} function(s), "
            f"{loops} loop(s), and {branches} conditional branch(es)."
        )
   
    # Framework context
    if fws:
        parts.append(f"It makes use of {', '.join(fws[:3])}.")
    elif imports:
        parts.append(f"It imports {', '.join(imports[:4])}.")

    # Control flow
    if loops > 0 and branches > 0:
        parts.append(
            f"The code includes {loops} loop{'s' if loops > 1 else ''} and "
            f"{branches} conditional branch{'es' if branches > 1 else ''} "
            f"for flow control."
        )
    elif loops > 0:
        parts.append(f"It uses {loops} loop{'s' if loops > 1 else ''} to iterate over data.")
    elif branches > 0:
        parts.append(f"It contains {branches} conditional check{'s' if branches > 1 else ''} for decision logic.")

    # Notable patterns
    if "recursion" in patterns:
        fn = next((f for f in funcs if text.count(f+"(") > 1), funcs[0] if funcs else "a function")
        parts.append(f"It uses recursion in '{fn}' to solve a problem by calling itself.")
    if "sorting algorithm" in patterns:
        parts.append("A sorting algorithm is implemented or invoked to order data.")
    if "search algorithm" in patterns:
        parts.append("A search algorithm is used to locate elements efficiently.")
    if "database operations" in patterns:
        parts.append("It interacts with a database to store or retrieve records.")
    if "HTTP/API calls" in patterns:
        parts.append("External APIs or HTTP endpoints are called to fetch or send data.")
    if "file handling" in patterns:
        parts.append("Files are read from or written to disk.")
    if "exception handling" in patterns:
        parts.append("Error handling is implemented to manage exceptions gracefully.")
    if "asynchronous/concurrent execution" in patterns:
        parts.append("Asynchronous or concurrent execution patterns are used for efficiency.")
    if "functional programming patterns" in patterns:
        parts.append("Functional programming techniques such as map, filter, or lambda are applied.")
   
    # ─────────────────────────────────────────────────────
    # SMART PURPOSE DETECTION
    # ─────────────────────────────────────────────────────

    # Python / Generic Algorithms
    if any("fibonacci" in f.lower() for f in funcs):
        return "This program generates Fibonacci numbers using recursion and prints the sequence."

    if any("factorial" in f.lower() for f in funcs):
        return "This program calculates the factorial of a number."

    if any("prime" in f.lower() for f in funcs):
        return "This program checks whether a number is prime."

    if any("palindrome" in f.lower() for f in funcs):
        return "This program determines whether the given input is a palindrome."

    if any("sort" in f.lower() for f in funcs):
        return "This program sorts data into a specific order."

    if any("search" in f.lower() for f in funcs):
        return "This program searches for a target value within a collection."

    if any("sum" in f.lower() for f in funcs):
        return "This program calculates the sum of values."

    if any("average" in f.lower() for f in funcs):
        return "This program computes the average of a dataset."

    if any("calculator" in f.lower() for f in funcs):
        return "This program performs arithmetic calculations."

    if any("encrypt" in f.lower() for f in funcs):
        return "This program encrypts information for secure storage or transmission."

    if any("decrypt" in f.lower() for f in funcs):
        return "This program decrypts previously encrypted information."

    # Java
    if lang == "Java":

        if "system.out.println" in t_low and loops > 0:
            return (
            "This Java program prints output to the console and processes values using loops."
            )

        if classes:
            return (
            "This Java program demonstrates object-oriented programming using classes and methods."
            )

        if "arraylist" in t_low:
            return (
            "This Java program stores and manages dynamic collections using ArrayList."
            )

        if "scanner" in t_low:
            return (
            "This Java program accepts user input from the console using Scanner."
            )

        if "jdbc" in t_low:
            return (
            "This Java program connects to a database and performs data operations."
            )

    # C
    if lang == "C":

        if "printf" in t_low:
            return (
            "This C program processes data and displays results using console output."
            )

        if "scanf" in t_low:
            return (
            "This C program accepts user input and processes the entered values."
            )

        if "malloc" in t_low:
            return (
            "This C program dynamically allocates memory and manages data structures."
            )

        if "struct" in t_low:
            return (
                "This C program uses structures to organize related data."
            )

    # C++
    if lang == "C++":

        if "cout" in t_low:
            return (
                "This C++ program processes information and displays results using standard output streams."
            )

        if classes:
            return (
                "This C++ program uses object-oriented programming concepts such as classes and objects."
            )

        if "vector" in t_low:
            return (
                "This C++ program stores and manipulates dynamic collections using vectors."
            )

        if "template" in t_low:
            return (
                "This C++ program uses templates for generic programming."
            )

    # SQL
    if lang == "SQL":

        if "SELECT" in text.upper():
            return (
            "This SQL query retrieves information from one or more database tables."
            )

        if "INSERT" in text.upper():
            return (
                "This SQL statement inserts new records into a database table."
             )

        if "UPDATE" in text.upper():
            return (
            "This SQL statement modifies existing records in a database table."
            )

        if "DELETE" in text.upper():
            return (
            "This SQL statement removes records from a database table."
            )

    # HTML
    if lang == "HTML":

        if "<form" in t_low:
            return (
            "This HTML document creates a form for collecting user input."
            )

        if "<table" in t_low:
            return (
            "This HTML document displays structured tabular information."
            )
    
        if "<nav" in t_low:
            return (
            "This HTML document contains navigation elements for website navigation."
            )

        return (
            "This HTML document defines the structure and content of a web page."
        )



    # JavaScript
    if lang == "JavaScript":

        if "fetch(" in text or "axios" in t_low:
            return (
                "This JavaScript program communicates with external APIs to retrieve or send data."
            )

        if "document." in text:
            return (
                "This JavaScript code manipulates webpage elements through the DOM."
            )

        if "addEventListener" in text:
            return (
                "This JavaScript code responds to user interactions through event listeners."
            )

    # Machine Learning
    if "machine learning" in patterns:

        if "linearregression" in t_low:
            return (
                "This program trains and evaluates a Linear Regression model."
            )
    
        if "randomforest" in t_low:
            return (
                "This program builds a Random Forest model for prediction."
            )

        if "tensorflow" in t_low or "keras" in t_low:
            return (
                "This program builds and trains a neural network."
            )

        if "torch" in t_low or "pytorch" in t_low:
            return (
                "This program uses a deep learning model implemented with PyTorch."
            )
    
        return (
            "This program trains or uses a machine learning model for prediction and analysis."
        )

    # Database
    if "database operations" in patterns:
        return (
            "This program interacts with a database to retrieve, store, update, or delete records."
        )

    # File Handling
    if "file handling" in patterns:
        return (
            "This program reads data from files and/or writes output to files."
        )

    # HTTP/API
    if "HTTP/API calls" in patterns:
        return (
            "This program communicates with external APIs or web services."
        )
    
    return " ".join(parts)


def _generate_code_smart_notes(info: dict, text: str) -> list:
    """
    Generates smart bullet-point notes for code.
    Returns a list of strings (without bullet prefix).
    """
    notes = []
    lang  = info["lang"]
    t_low = text.lower()

    for p in info["patterns"][:5]:
        notes.append(f"Demonstrates {p}")

    if info["classes"]:
        notes.append(f"Defines {len(info['classes'])} class(es): {', '.join(info['classes'][:3])}")
    if info["functions"]:
        notes.append(f"Contains {len(info['functions'])} function(s): {', '.join(info['functions'][:4])}")
    if info["loops"] > 0:
        notes.append(f"Uses {info['loops']} loop(s) for iteration")
    if info["branches"] > 0:
        notes.append(f"Has {info['branches']} conditional branch(es)")
    if info["frameworks"]:
        notes.append(f"Libraries/frameworks: {', '.join(info['frameworks'][:3])}")

    # Language-specific notes — 7 supported languages only
    if lang == "Python":
        if "if __name__" in t_low: notes.append("Script can run standalone or be imported as a module")
        if re.search(r"\bwith\s+open\b", text): notes.append("Uses context manager for safe file handling")
        if re.search(r"@\w+", text): notes.append("Uses decorators for extending function behaviour")
        if re.search(r"\byield\b", text): notes.append("Uses generators (yield) for memory-efficient iteration")
    elif lang == "Java":
        if re.search(r"@Override", text): notes.append("Overrides parent class methods using @Override")
        if re.search(r"interface\s+\w+", text): notes.append("Defines or implements an interface")
    elif lang == "JavaScript":
        if re.search(r"async|await|Promise", text): notes.append("Uses async/await for non-blocking operations")
        if re.search(r"document\.|getElementById|querySelector", text): notes.append("Manipulates the DOM directly")
    elif lang == "C++":
        if re.search(r"\btemplate\s*<", text): notes.append("Uses templates for generic programming")
        if re.search(r"\bnew\b|\bdelete\b", text): notes.append("Manages heap memory manually with new/delete")
    elif lang == "C":
        if re.search(r"\bmalloc\b|\bcalloc\b", text): notes.append("Manages heap memory with malloc/free")
        if re.search(r"\bpointer\b|\*\w+", text): notes.append("Uses pointer arithmetic for direct memory access")
    elif lang == "SQL":
        if re.search(r"\bJOIN\b", text, re.I): notes.append("Joins multiple tables to combine related data")
        if re.search(r"\bGROUP\s+BY\b", text, re.I): notes.append("Aggregates rows using GROUP BY")
        if re.search(r"\bINDEX\b", text, re.I): notes.append("Defines an index for query optimisation")
    elif lang == "HTML":
        if re.search(r"<form", text, re.I): notes.append("Contains a form for user input")
        if re.search(r"<script", text, re.I): notes.append("Embeds JavaScript for interactivity")

    return list(dict.fromkeys(notes))[:7]


def _explain_code(text: str):
    """
    Main code explanation entry point.
    Returns (explanation_string, language) or (None, None) if not code.

    Output format:
    SUMMARY
    [natural language summary]

    💻 Code Analysis
    Language: X  |  Size: N lines  |  Functions: N  |  Classes: N  |  Loops: N  |  Branches: N

    WHAT THIS CODE DOES
    [detailed natural language explanation]
    """
    lang = _detect_code(text)
    if not lang:
        return None, None

    info    = _analyze_code_structure(text, lang)
    narrative = _generate_code_narrative(info, text)

    lines    = info["total_lines"]
    funcs    = info["functions"]
    classes  = info["classes"]
    loops    = info["loops"]
    branches = info["branches"]
    patterns = info["patterns"]
    fws      = info["frameworks"]
    imports  = info["imports"]
    size     = info["size_label"]

    # ── SUMMARY block (maps to AI Summary output) ────────────
    summary_line = (
        f"This {lang} program "
        f"contains {len(funcs)} function(s), "
        f"{loops} loop(s), and "
        f"{branches} conditional branch(es)."
    )

    # ── Analysis block ────────────────────────────────────────
    func_str   = ", ".join(funcs[:4]) if funcs   else "0"
    class_str  = ", ".join(classes[:3]) if classes else "0"

    analysis = (
        f"Language: {lang}  |  Size: {lines} lines ({size})  |  "
        f"Functions: {func_str}  |  Classes: {class_str}  |  "
        f"Loops: {loops}  |  Branches: {branches}"
    )
    if imports:
        analysis += f"\nImports: {', '.join(imports[:6])}"
    if fws:
        analysis += f"\nFrameworks: {', '.join(fws[:4])}"
    if patterns:
        analysis += f"\nPatterns: {', '.join(patterns[:5])}"

    # ── Full output ───────────────────────────────────────────
    out_parts = [
        "SUMMARY",
        summary_line,
        "",
        "💻 Code Analysis",
        analysis,
        "",
        "WHAT THIS CODE DOES",
        narrative,
    ]

    return "\n".join(out_parts), lang


def _classify_doc(text, source=""):
    if source not in ("Audio","CSV/Excel"):
        lang = _detect_code(text)
        if lang: return f"Code / {lang}"
    t = text.lower()
    if any(w in t for w in ["abstract","methodology","hypothesis"]): return "Research Paper"
    if any(w in t for w in ["whereas","clause","provision","legal"]):  return "Legal Document"
    if any(w in t for w in ["revenue","profit","fiscal"]):             return "Financial Report"
    if any(w in t for w in ["patient","diagnosis","clinical"]):        return "Medical Document"
    if any(w in t for w in ["reported","journalist","breaking"]):      return "News Article"
    if source == "Audio":     return "Audio Transcript"
    if source == "CSV/Excel": return "Dataset / Tabular"
    if source == "PDF":       return "PDF Document"
    return "General Text"


# ──────────────────────────────────────────────────────────────
#  STEP 10 — KEYWORD EXTRACTION
# ──────────────────────────────────────────────────────────────
def _extract_kw(text, top_n=14):
    try:
        if not text or len(text.split()) < 5:
            return []
        vec   = TfidfVectorizer(stop_words="english", ngram_range=(1,2), max_features=600)
        tfidf = vec.fit_transform([text])
        names  = vec.get_feature_names_out()
        scores = tfidf.toarray()[0]
        scored = sorted(zip(names, scores), key=lambda x: x[1], reverse=True)
        return [k for k, v in scored[:top_n] if v > 0]
    except Exception:
        words    = re.findall(r"\b[a-zA-Z]{4,}\b", text.lower())
        stop     = {"this","that","with","from","have","been","will",
                    "they","their","about","there","which","would"}
        filtered = [w for w in words if w not in stop]
        return [w for w, _ in Counter(filtered).most_common(top_n)]


# ──────────────────────────────────────────────────────────────
#  STEP 9 — SMART NOTES (redesigned)
#  Extracts key facts, conclusions, action points, and concepts
#  from the English summary using semantic importance ranking.
# ──────────────────────────────────────────────────────────────
def _smart_notes(eng_sum: str, doc_type: str, kws: list) -> str:
    """
    Generates intelligent Smart Notes from the English summary.
    Uses semantic ranking (all-mpnet-base-v2) to select the most
    important sentences as bullet points — not just the first N.

    Categories extracted:
    - Key facts & findings
    - Important concepts
    - Main ideas
    - Conclusions
    - Action points (where applicable)
    """
    try:
        def get_sentences(text):
            try:    return nltk.sent_tokenize(text)
            except: return re.split(r"(?<=[.!?])\s+", text)

        sents = [s.strip() for s in get_sentences(eng_sum) if s.strip()]

        # ── Filter for quality ────────────────────────────────
        # Keep sentences that are substantial (25–280 chars)
        quality_sents = [s for s in sents if 25 < len(s) < 280]

        if not quality_sents:
            quality_sents = [s for s in sents if len(s.strip()) > 10]

        if not quality_sents:
            return "• " + eng_sum[:200].strip()

        # ── Semantic selection for smart notes ───────────────
        # Select top 5 most important sentences using all-mpnet-base-v2
        num_bullets = min(5, len(quality_sents))
        if len(quality_sents) > num_bullets:
            top_sents = _semantic_rank(quality_sents, top_k=num_bullets)
        else:
            top_sents = quality_sents

        # ── Categorise bullets by type ────────────────────────
        # Try to detect sentence roles for richer notes
        def _label_sentence(s: str) -> str:
            sl = s.lower()
            if any(w in sl for w in ["conclud","therefore","thus","in summary","result","finding"]):
                return "📌"
            if any(w in sl for w in ["should","must","recommend","need to","important to","require"]):
                return "⚡"
            if any(w in sl for w in ["show","demonstrate","reveal","found","indicate","suggest"]):
                return "🔬"
            if any(w in sl for w in ["increase","decrease","grow","rise","fall","percent","%"]):
                return "📊"
            return "•"

        pts = [f"{_label_sentence(s)} {s}" for s in top_sents]
        body = "\n".join(pts)

        # ── Quick Ideas by document type ─────────────────────
        qi_map = {
            "Research":  ["📌 Main Finding","🔬 Evidence","🔮 Future Scope","📊 Methodology"],
            "Legal":     ["⚖️ Key Obligation","📋 Key Clause","🚨 Risk","📝 Action"],
            "Financial": ["💰 Key Metric","📈 Growth","⚠️ Risk","🎯 Target"],
            "Code":      ["💻 Purpose","🔧 Core Logic","📦 Libraries","🧪 Test"],
            "Medical":   ["🏥 Finding","💊 Treatment","⚕️ Indicator","🔬 Lab"],
            "News":      ["📰 Event","🌍 Impact","🕐 Timeline","👥 Person"],
            "Audio":     ["🎙️ Statement","💬 Topic","📢 Quote","🗓️ Context"],
            "Dataset":   ["📊 Shape","📉 Missing","📈 Trend","💡 Insight"],
        }
        doc_key = next((k for k in qi_map if k.lower() in doc_type.lower()), None)
        quick   = qi_map.get(doc_key, ["✨ Key Insight","📝 Remember","🎯 Takeaway","💡 Note"])

        out = (f"🧠 SMART NOTES\n{'─'*28}\n{body}\n\n"
               f"💡 QUICK IDEAS\n{'  '.join(quick[:4])}")

        if kws:
            out += f"\n\n🔑 KEY TERMS\n{'  ·  '.join(kws[:5])}"

        return out

    except Exception as e:
        logger.warning(f"Smart notes error: {e}")
        return "Smart notes could not be generated."


# ──────────────────────────────────────────────────────────────
#  STEP 8 — TITLE GENERATION
# ──────────────────────────────────────────────────────────────
def _gen_title(orig, eng_sum, kws, doc_type):
    bad_starts = (
        "page","chapter","section","figure","table",
        "introduction","overview","summary","this document",
        "the following","in this","it is","we can",
    )
    if orig:
        for line in orig.strip().split("\n")[:25]:
            line  = line.strip().rstrip(".,:;-=")
            words = line.split()
            if not (3 <= len(words) <= 14): continue
            if line.lower().startswith(bad_starts): continue
            if re.match(r"^\d+[\.\)]\s*", line):   continue
            if len(line) < 4:                       continue
            return line[:80]
    try:
        sents = nltk.sent_tokenize(eng_sum)
    except Exception:
        sents = re.split(r"(?<=[.!?])\s+", eng_sum)
    if sents:
        first = sents[0].strip().rstrip(".!?,;:")
        words = first.split()
        if 5 <= len(words) <= 12:  return first[:80]
        if len(words) > 12:        return " ".join(words[:10]) + "…"
    if kws:
        return " | ".join(k.title() for k in kws[:3])
    return doc_type if doc_type else "AI Summary"


# ──────────────────────────────────────────────────────────────
#  INPUT EXTRACTORS
# ──────────────────────────────────────────────────────────────
def _extract_pdf(path):
    try:
        pages = []
        with fitz.open(path) as doc:
            for p in doc:
                text = p.get_text("text")
                if text and text.strip():
                    pages.append(text.strip())
        result = "\n".join(pages)
        if not result.strip():
            return "PDF Error: no extractable text (scanned/image-based)."
        logger.info(f"PDF: {_wc(result)} words, {len(pages)} pages")
        return result
    except Exception as e:
        return f"PDF Error: {e}"


def _transcribe(path):
    try:
        asr    = _get_whisper()
        result = asr(path, generate_kwargs={"language":"en","task":"transcribe"})
        text   = result.get("text","").strip()
        if not text: return None, "Empty transcription."
        return text, f"Transcribed: {_wc(text)} words"
    except Exception as e:
        logger.warning(f"Whisper primary failed: {e}")
        try:
            result = _get_whisper()(path, generate_kwargs={"language":"en","task":"transcribe"})
            text   = result.get("text","").strip()
            if text: return text, f"Transcribed: {_wc(text)} words"
        except Exception as e2:
            logger.error(f"Whisper fallback failed: {e2}")
        return None, f"Audio error: {e}"


def _analyze_data(path):
    try:
        p  = str(path).lower()
        df = (pd.read_csv(path, encoding="utf-8", on_bad_lines="skip")
              if p.endswith(".csv") else pd.read_excel(path))
        rows, cols = df.shape
        num_cols   = df.select_dtypes(include="number").columns.tolist()
        txt_cols   = df.select_dtypes(include="object").columns.tolist()
        missing    = int(df.isnull().sum().sum())
        stat_lines = [
            f"Shape: {rows:,} rows × {cols} columns",
            f"Columns: {', '.join(list(df.columns)[:10])}{'…' if cols>10 else ''}",
            f"Numeric ({len(num_cols)}): {', '.join(num_cols[:6]) or 'none'}",
            f"Text ({len(txt_cols)}): {', '.join(txt_cols[:6]) or 'none'}",
            f"Missing: {missing} ({round(missing/(rows*cols)*100,1) if rows*cols else 0}%)",
        ]
        insight_lines = []
        for col in num_cols[:5]:
            s = df[col].describe()
            mn, mx_, me, std = (round(s[k],2) for k in ["min","max","mean","std"])
            stat_lines.append(f"  {col}: min={mn}  max={mx_}  mean={me}  std={std}")
            vals = df[col].dropna().tolist()
            if len(vals) >= 6:
                t = len(vals)//3
                a_s = sum(vals[:t])/t; a_e = sum(vals[-t:])/t
                if a_s != 0:
                    pct = round((a_e-a_s)/abs(a_s)*100, 1)
                    if pct > 20:    insight_lines.append(f"📈 {col} trending UP +{pct}%")
                    elif pct < -20: insight_lines.append(f"📉 {col} trending DOWN {pct}%")
        for col in txt_cols[:3]:
            vc = df[col].value_counts()
            if len(vc): insight_lines.append(
                f"🏷️ Top '{col}': {vc.index[0]} ({round(vc.iloc[0]/rows*100,1)}%)")
        show    = list(df.columns)[:7]
        hdr     = "| " + " | ".join(str(c)[:14] for c in show) + " |"
        sep     = "| " + " | ".join(["---"]*len(show)) + " |"
        rows_md = [hdr,sep] + ["| "+" | ".join(str(v)[:14] for v in row.values)+" |"
                                for _, row in df[show].head(5).iterrows()]
        stat_text  = " ".join(stat_lines+insight_lines)
        insights   = "\n".join(insight_lines) if insight_lines else "No strong trends."
        report     = "\n".join(stat_lines) + "\n\n💡 INSIGHTS:\n" + insights
        return report, stat_text, "\n".join(rows_md), insights
    except Exception as e:
        return f"Data error: {e}", "", "", ""


# ──────────────────────────────────────────────────────────────
#  TTS — +10 dB
# ──────────────────────────────────────────────────────────────
def _tts(text, lang_name):
    try:
        from gtts import gTTS
        code       = LANGUAGES.get(lang_name, "en")
        gtts_code  = GTTS_CODES.get(code, code.split("-")[0])
        tts = gTTS(text=text[:3000], lang=gtts_code, slow=False)
        raw = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False).name
        tts.save(raw)
        try:
            from pydub import AudioSegment
            louder = AudioSegment.from_mp3(raw) + 10
            out    = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False).name
            louder.export(out, format="mp3", bitrate="128k")
            return out
        except: return raw
    except Exception as e:
        logger.warning(f"TTS: {e}"); return None


# ──────────────────────────────────────────────────────────────
#  PDF EXPORT
# ──────────────────────────────────────────────────────────────
def _export_pdf(orig, summ, title, notes, kws, det_lang, out_lang,
                translated, model, tone, sb, sa, ratio, ts, trans_text=""):
    pdf = FPDF(); pdf.set_auto_page_break(auto=True, margin=14); pdf.add_page()

    def div():
        pdf.set_draw_color(0,212,255); pdf.set_line_width(0.5)
        pdf.line(10,pdf.get_y(),200,pdf.get_y()); pdf.ln(4)

    def sec(t, color=(0,150,200)):
        pdf.ln(3); pdf.set_font("Helvetica","B",11); pdf.set_text_color(*color)
        safe = _strip_emojis(_clean(t)).encode("latin-1","replace").decode("latin-1")
        pdf.cell(0,8,safe,new_x=XPos.LMARGIN,new_y=YPos.NEXT)
        pdf.set_text_color(30,30,30); pdf.set_font("Helvetica","",10)

    def kv(k, v):
        pdf.set_font("Helvetica","B",10); pdf.set_text_color(50,80,120)
        ks = _strip_emojis(_clean(k)).encode("latin-1","replace").decode("latin-1")
        vs = _strip_emojis(_clean(str(v))).encode("latin-1","replace").decode("latin-1")
        pdf.cell(52,6,ks+":",new_x=XPos.RIGHT,new_y=YPos.TOP)
        pdf.set_font("Helvetica","",10); pdf.set_text_color(20,20,50)
        pdf.cell(0,6,vs,new_x=XPos.LMARGIN,new_y=YPos.NEXT)

    def blk(text, fill=(235,248,255), tcol=(10,30,80), fs=10):
        safe = _strip_emojis(_clean(text)).encode("latin-1","replace").decode("latin-1")
        pdf.set_font("Helvetica","",fs); pdf.set_text_color(*tcol)
        pdf.set_fill_color(*fill); pdf.multi_cell(0,5.5,safe,fill=True)

    pdf.set_font("Helvetica","B",22); pdf.set_text_color(0,100,160)
    pdf.cell(0,16,"INSTASUMMARY REPORT",new_x=XPos.LMARGIN,new_y=YPos.NEXT,align="C")
    pdf.set_font("Helvetica","I",10); pdf.set_text_color(100,120,140)
    pdf.cell(0,6,_clean(f"Generated: {ts}"),new_x=XPos.LMARGIN,new_y=YPos.NEXT,align="C")
    pdf.ln(4); pdf.set_draw_color(0,212,255); pdf.set_line_width(1.2)
    pdf.line(10,pdf.get_y(),200,pdf.get_y()); pdf.ln(8)

    if title:
        sec("Document Title", color=(0,100,160))
        pdf.set_font("Helvetica","B",13); pdf.set_text_color(0,100,160)
        safe_title = _strip_emojis(_clean(title)).encode("latin-1","replace").decode("latin-1")
        pdf.multi_cell(0,7,safe_title); pdf.ln(2); div()

    sec("Report Metadata")
    kv("AI Model",model); kv("Tone",tone)
    kv("Detected Language",det_lang); kv("Output Language",out_lang)
    kv("Translated","Yes" if translated else "No"); pdf.ln(2); div()

    sec("Statistics")
    hdrs  = ["","Words","Sentences","Read Time"]
    rdata = [["Original",str(sb["words"]),str(sb["sents"]),sb["rt"]],
             ["Summary", str(sa["words"]),str(sa["sents"]),sa["rt"]]]
    cw    = [36,34,34,34]
    pdf.set_fill_color(0,100,160); pdf.set_text_color(255,255,255); pdf.set_font("Helvetica","B",9)
    for i,h in enumerate(hdrs): pdf.cell(cw[i],6,_clean(h),border=1,fill=True)
    pdf.ln()
    for ri, row in enumerate(rdata):
        pdf.set_fill_color(*(235,248,255) if ri==0 else (235,255,245))
        pdf.set_text_color(20,20,50); pdf.set_font("Helvetica","",9)
        for i,v in enumerate(row): pdf.cell(cw[i],6,_clean(v),border=1,fill=True)
        pdf.ln()
    pdf.ln(2); pdf.set_font("Helvetica","B",9); pdf.set_text_color(0,150,200)
    pdf.cell(0,6,f"Compression: {ratio}%",new_x=XPos.LMARGIN,new_y=YPos.NEXT)
    pdf.ln(2); div()

    sec("Keywords"); pdf.set_font("Helvetica","",10); pdf.set_text_color(30,60,120)
    safe_kw = _strip_emojis(_clean("  ·  ".join(kws))).encode("latin-1","replace").decode("latin-1")
    pdf.multi_cell(0,6,safe_kw); pdf.ln(2); div()

    sec("AI Summary", color=(0,100,160))
    blk(summ, fill=(235,248,255), tcol=(10,30,80)); pdf.ln(2); div()

    if notes and notes.strip():
        sec("Smart Notes", color=(20,100,40))
        blk(notes, fill=(240,255,240), tcol=(10,60,20)); pdf.ln(2); div()

    if trans_text and trans_text.strip():
        sec("Translation", color=(80,30,160))
        blk(trans_text, fill=(245,240,255), tcol=(40,10,100)); pdf.ln(2); div()

    if pdf.get_y() > 220: pdf.add_page()
    sec("Original Source Text", color=(20,100,40))
    blk(orig[:10000]+("\n…[truncated]" if len(orig)>10000 else ""),
        fill=(240,255,240), tcol=(10,50,20), fs=9)

    pdf.ln(6); pdf.set_font("Helvetica","I",8); pdf.set_text_color(150,165,180)
    pdf.cell(0,5,"InstaSummary | Powered by AI | Hugging Face Spaces",
             new_x=XPos.LMARGIN,new_y=YPos.NEXT,align="C")

    out = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False).name
    pdf.output(out); return out


# ──────────────────────────────────────────────────────────────
#  DOCX EXPORT
# ──────────────────────────────────────────────────────────────
def _export_docx(orig, summ, title, notes, kws, det_lang, out_lang,
                 sb, sa, ratio, ts, trans_text=""):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        doc = Document()
        hd  = doc.add_heading(title or "InstaSummary Report", 0)
        hd.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(f"Generated: {ts}  |  {det_lang} → {out_lang}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("")
        doc.add_heading("AI Generated Summary", 1); doc.add_paragraph(summ); doc.add_paragraph("")
        if notes and notes.strip():
            doc.add_heading("Smart Notes", 1); doc.add_paragraph(notes); doc.add_paragraph("")
        if trans_text and trans_text.strip():
            doc.add_heading(f"Translation ({out_lang})", 1)
            doc.add_paragraph(trans_text); doc.add_paragraph("")
        doc.add_heading("Keywords", 1)
        doc.add_paragraph("  ·  ".join(kws)); doc.add_paragraph("")
        doc.add_heading("Statistics", 1)
        tbl = doc.add_table(rows=3, cols=4); tbl.style = "Table Grid"
        for i,h in enumerate(["","Words","Sentences","Read Time"]): tbl.rows[0].cells[i].text = h
        for i,v in enumerate(["Original",str(sb["words"]),str(sb["sents"]),sb["rt"]]): tbl.rows[1].cells[i].text = v
        for i,v in enumerate(["Summary", str(sa["words"]),str(sa["sents"]),sa["rt"]]): tbl.rows[2].cells[i].text = v
        doc.add_paragraph(f"Compression: {ratio}%"); doc.add_paragraph("")
        doc.add_heading("Original Source Text", 1)
        doc.add_paragraph(orig[:8000]+("…[truncated]" if len(orig)>8000 else ""))
        doc.add_paragraph(""); doc.add_paragraph("Built with InstaSummary | Powered by AI")
        out = tempfile.NamedTemporaryFile(suffix=".docx", delete=False).name
        doc.save(out); return out
    except Exception as e:
        logger.warning(f"DOCX: {e}")
        try:
            f = tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w", encoding="utf-8")
            f.write(f"INSTASUMMARY REPORT\n{'='*40}\n{ts}\n"
                    f"TITLE:\n{title}\nSUMMARY:\n{summ}\n"
                    f"KEYWORDS:\n{', '.join(kws)}\nORIGINAL:\n{orig[:6000]}\n")
            f.close(); return f.name
        except: return None


# ──────────────────────────────────────────────────────────────
#  HTML BUILDERS
# ──────────────────────────────────────────────────────────────
def _status(msg, kind="ok"):
    cfg = {"ok":   ("#10b981","rgba(16,185,129,.09)","rgba(16,185,129,.22)","✅"),
           "err":  ("#ef4444","rgba(239,68,68,.09)",  "rgba(239,68,68,.22)",  "⚠️"),
           "info": ("#00d4ff","rgba(0,212,255,.09)",  "rgba(0,212,255,.22)",  "⚡")}
    c, bg, br, ico = cfg.get(kind, cfg["info"])
    return (f'<div style="display:flex;align-items:center;gap:9px;background:{bg};'
            f'border:1px solid {br};border-radius:8px;padding:9px 14px;'
            f'color:{c};font-size:.85rem;font-weight:500;font-family:Inter,sans-serif">'
            f'{ico} {msg}</div>')

def _info_bar(lang_info, dtype):
    if not lang_info and not dtype: return ""
    return (f'<div style="display:flex;gap:10px;flex-wrap:wrap;margin:8px 0">'
            f'<span style="background:rgba(0,212,255,.08);border:1px solid rgba(0,212,255,.18);'
            f'color:#67e8f9;border-radius:7px;padding:4px 12px;font-size:.76rem;font-weight:500">'
            f'🌐 {lang_info}</span>'
            f'<span style="background:rgba(139,92,246,.08);border:1px solid rgba(139,92,246,.18);'
            f'color:#c4b5fd;border-radius:7px;padding:4px 12px;font-size:.76rem;font-weight:500">'
            f'📄 {dtype}</span></div>')

def _title_card(t):
    if not t: return ""
    return (f'<div style="background:linear-gradient(135deg,rgba(0,212,255,.06),rgba(56,189,248,.04));'
            f'border:1px solid rgba(0,212,255,.18);border-radius:10px;padding:11px 15px;margin-bottom:8px">'
            f'<div style="color:#6b7280;font-size:.6rem;letter-spacing:2px;text-transform:uppercase;'
            f'font-weight:700;margin-bottom:3px">✨ DOCUMENT TITLE</div>'
            f'<div style="color:#e0f2fe;font-size:.96rem;font-weight:700;line-height:1.4;'
            f'font-family:Poppins,sans-serif">{t}</div></div>')

def _kw_html(kws):
    if not kws:
        return '<p style="color:#374151;padding:8px;font-size:.83rem;margin:0">Keywords appear after summarizing.</p>'
    pills = "".join(
        f'<span style="background:rgba(0,212,255,.1);border:1px solid rgba(0,212,255,.25);'
        f'color:#67e8f9;border-radius:100px;padding:4px 12px;font-size:.74rem;'
        f'font-weight:500;display:inline-block;margin:3px 3px">{k}</span>'
        for k in kws
    )
    return f'<div style="padding:4px 2px">{pills}</div>'

def _hist_html(history):
    if not history:
        return ('<div style="text-align:center;padding:26px;color:#374151">'
                '<div style="font-size:1.5rem;margin-bottom:6px">📭</div>'
                '<div style="font-size:.85rem">No history yet.</div></div>')
    cards = ""
    for i, h in enumerate(history):
        cards += (f'<div style="display:flex;align-items:center;justify-content:space-between;'
                  f'background:rgba(6,21,46,.65);border:1px solid rgba(0,212,255,.1);'
                  f'border-radius:9px;padding:10px 14px;margin-bottom:6px;transition:border-color .2s" '
                  f'onmouseover="this.style.borderColor=\'rgba(0,212,255,.38)\'" '
                  f'onmouseout="this.style.borderColor=\'rgba(0,212,255,.1)\'">'
                  f'<div style="flex:1;margin-right:10px">'
                  f'<div style="color:#e0f2fe;font-weight:600;font-size:.82rem;margin-bottom:2px">{h["title"]}</div>'
                  f'<div style="color:#4b5563;font-size:.70rem">{h["time"]} · {h["type"]}</div></div>'
                  f'<span style="background:rgba(0,212,255,.08);border:1px solid rgba(0,212,255,.18);'
                  f'color:#22d3ee;border-radius:5px;padding:2px 9px;font-size:.68rem;font-weight:600">#{i+1}</span></div>')
    return cards

def _ty_html():
    return ('<div style="text-align:center;padding:16px;'
            'background:linear-gradient(135deg,rgba(0,212,255,.05),rgba(20,184,166,.04));'
            'border:1px solid rgba(0,212,255,.12);border-radius:12px;margin-top:7px">'
            '<div style="font-size:1.15rem;margin-bottom:4px">🎉</div>'
            '<div style="color:#e0f2fe;font-weight:700;font-size:.86rem;margin-bottom:2px">Summary Generated Successfully!</div>'
            '<div style="color:#64748b;font-size:.77rem">Thank you for using InstaSummary</div></div>')

def _err(msg):
    return (
        _status(msg, "err"),  # 0
        "",                   # 1  info_html
        "",                   # 2  title_html
        "",                   # 3  summary_out
        None,                 # 4  voice_out
        "","","",             # 5,6,7
        "","","",             # 8,9,10
        "",                   # 11 ratio_out
        _kw_html([]),         # 12 kw_html
        "","",                # 13,14 compare
        "*No data.*","",      # 15,16 data
        None, None,           # 17,18 pdf,docx
        "",                   # 19 ty_html
        "",                   # 20 notes_out
    )


# ──────────────────────────────────────────────────────────────
#  MAIN PIPELINE — 21 outputs
#  Integrated BERT + DistilBART Smart Summarization Engine
# ──────────────────────────────────────────────────────────────
def _pipeline(raw_text, source, model_choice, length, tone, out_lang, fmt, csv_md=""):
    if not raw_text or _wc(raw_text) < 15:
        return _err("Input too short — need at least 15 words.")
    if not _model_ready():
        return _err("Model loading — wait ~30 seconds and retry.")

    is_data  = source in ("Audio","CSV/Excel")
    src_code = _dl(raw_text)
    src_name = _cn(src_code)
    out_code = LANGUAGES.get(out_lang, "en")

    # ── Step 3: Language detection & translation to English ──
    if src_code != "en":
        if is_data and len(raw_text) > 8000:
            eng = (_translate(raw_text[:8000], "en") + " " +
                   _translate(raw_text[8000:16000], "en")).strip()
        else:
            eng = _translate(raw_text, "en")
        translated = True
    else:
        eng        = raw_text
        translated = False

    # ── Step 2: Enhanced text cleaning ───────────────────────
    eng = _preprocess(eng)

    wb   = _wc(eng); sb_n = _sc(eng); rb = _rt(wb)
    logger.info(f"Summarising {wb} words | BERT+MPNet | → {out_lang}")

    mx, mn = LENGTH_MAP.get(length, (150, 100))
    if out_code != "en":
        mx += LANG_BONUS
        mn += LANG_BONUS // 2

    # ── Code detection ────────────────────────────────────────
    code_exp, code_lang = _explain_code(eng) if not is_data else (None, None)

    if code_exp:
        # Code path: bypass _fmt_summary (structured output must not be truncated)
        eng_sum = code_exp
        # Generate code-specific smart notes
        _code_info = _analyze_code_structure(eng, code_lang)
        _code_note_lines = _generate_code_smart_notes(_code_info, eng)
        code_smart_notes = (
            "🧠 SMART NOTES\n" + "─"*28 + "\n" +
            "\n".join(f"• {n}" for n in _code_note_lines) +
            "\n\n💡 QUICK IDEAS\n💻 Purpose  🔧 Core Logic  📦 Libraries  🧪 Test"
        )
        if _code_info.get("imports"):
            code_smart_notes += f"\n\n🔑 KEY TERMS\n{'  ·  '.join(_code_info['imports'][:5])}"
    else:
        src_text = eng
        # For large data sources, trim intelligently
        if is_data and _wc(eng) > 4000:
            words    = eng.split(); n = len(words); third = n // 3
            src_text = " ".join(
                words[:min(1400, third)] + ["[…]"] +
                words[n//2-400:n//2+400] + ["[…]"] +
                words[max(0, n-1000):]
            )
        # ── Steps 4–6: Semantic ranking + hierarchical summarization ──
        eng_sum = _hierarchical(src_text, mx, mn)

    # ── Step 7: Format & tone post-processing ─────────────────
    # Code output is already structured — skip _fmt_summary for it
    if not code_exp:
        eng_sum = _fmt_summary(_apply_tone(eng_sum, tone), fmt, length)

    # ── Step 3 (output): Translate back to target language ───
    final = _translate(eng_sum, out_code, chunk_size=1500) if out_code != "en" else eng_sum

    wa   = _wc(final); sa_n = _sc(final); ra = _rt(wa)
    ratio    = _compress(eng, final)
    kws      = _extract_kw(eng)
    doc_type = _classify_doc(eng, source)
    title    = _gen_title(eng, eng_sum, kws, doc_type)
    # Use code-specific notes if code was detected, otherwise prose smart notes
    notes    = code_smart_notes if code_exp else _smart_notes(eng_sum, doc_type, kws)
    orig_preview = eng[:700] + ("…" if len(eng) > 700 else "")
    ts  = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
    sb  = {"words": wb,  "sents": sb_n, "rt": rb}
    sa  = {"words": wa,  "sents": sa_n, "rt": ra}

    voice     = _tts(final, out_lang)
    lang_info = (f"{source} · {src_name} → {out_lang}" +
                 (" (translated)" if translated else ""))

    pdf_p  = _export_pdf(eng, final, title, notes, kws, src_name, out_lang,
                         translated, "BERT+MPNet", tone, sb, sa, ratio, ts)
    docx_p = _export_docx(eng, final, title, notes, kws, src_name, out_lang,
                          sb, sa, ratio, ts)

    logger.info("Pipeline complete")
    return (
        _status(f"Done! {source} → {out_lang}", "ok"),  # 0
        _info_bar(lang_info, doc_type),                  # 1
        _title_card(title),                               # 2
        final,                                            # 3
        voice,                                            # 4
        str(wb), str(sb_n), rb,                          # 5,6,7
        str(wa), str(sa_n), ra,                          # 8,9,10
        f"{ratio}%",                                      # 11
        _kw_html(kws),                                    # 12
        orig_preview, final,                              # 13,14
        csv_md,                                           # 15
        "",                                               # 16
        pdf_p, docx_p,                                    # 17,18
        _ty_html(),                                       # 19
        notes,                                            # 20
    )


# ──────────────────────────────────────────────────────────────
#  HANDLERS
# ──────────────────────────────────────────────────────────────
def process_text(t, m, le, to, la, fm):
    try:
        if t and str(t).strip(): return _pipeline(str(t).strip(),"Text",m,le,to,la,fm)
        return _err("Please paste text to summarize.")
    except Exception as e: logger.exception("text"); return _err(f"Error: {e}")

def process_pdf(p, m, le, to, la, fm):
    try:
        path = _sp(p)
        if not path: return _err("Please upload a PDF.")
        raw = _extract_pdf(path)
        if raw.startswith("PDF Error"): return _err(raw)
        return _pipeline(raw,"PDF",m,le,to,la,fm)
    except Exception as e: logger.exception("pdf"); return _err(f"PDF Error: {e}")

def process_audio(a, m, le, to, la, fm):
    try:
        if a is None: return _err("Please upload or record audio.")
        path = _sp(a)
        if not path: return _err("Could not read audio.")
        text, status = _transcribe(path)
        if text is None: return _err(status)
        n = _wc(text)
        if n > 6000:   le = "Long (200-280)"
        elif n > 3000: le = "Medium (100-150)"
        return _pipeline(text,"Audio",m,le,to,la,fm)
    except Exception as e: logger.exception("audio"); return _err(f"Audio error: {e}")

def process_csv(c, m, le, to, la, fm):
    try:
        if c is None: return _err("Please upload a CSV or Excel file.")
        path = _sp(c)
        if not path: return _err("Could not read file.")
        report, stat_text, preview, insights = _analyze_data(path)
        if not stat_text: return _err(str(report))
        result = list(_pipeline(stat_text,"CSV/Excel",m,le,to,la,fm,preview))
        result[3]  = f"📊 DATA ANALYSIS:\n{report}\n\n📝 AI SUMMARY:\n{result[3]}"
        result[15] = preview
        result[16] = insights
        try:
            kws_t = _extract_kw(stat_text)
            sb_t  = {"words": _wc(stat_text), "sents": _sc(stat_text), "rt": ""}
            sa_t  = {"words": _wc(result[3]), "sents": _sc(result[3]),  "rt": ""}
            ts    = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
            notes_t = _smart_notes(result[3], "Dataset / Tabular", kws_t)
            result[17] = _export_pdf(stat_text, result[3], "Data Analysis", notes_t, kws_t,
                                     "Auto", la, False, "BERT+MPNet", "Professional",
                                     sb_t, sa_t, 0, ts)
            result[18] = _export_docx(stat_text, result[3], "Data Analysis", notes_t, kws_t,
                                      "Auto", la, sb_t, sa_t, 0, ts)
        except: pass
        return tuple(result)
    except Exception as e: logger.exception("csv"); return _err(f"Data error: {e}")


def translate_and_voice(summary, trans_lang):
    if not summary or len(summary.strip()) < 5: return "Generate a summary first.", None
    code = LANGUAGES.get(trans_lang, "en")
    if code == "en": return summary, _tts(summary, "English")
    try:
        translated = _translate(summary, code, chunk_size=1500)
        return translated, _tts(translated, trans_lang)
    except Exception as e:
        return f"Translation error: {e}", None

def update_history(summary, dtype, history):
    if not summary or len(summary.strip()) < 10:
        return history or [], _hist_html(history or [])
    title = summary.split("\n")[0][:70].strip()
    entry = {"title": title, "type": dtype or "General",
             "time": datetime.datetime.now().strftime("%H:%M · %b %d"), "summary": summary}
    new_hist = [entry] + list(history or [])[:9]
    _save_hist(new_hist)
    return new_hist, _hist_html(new_hist)

def fb_star(n):
    return f'<div style="text-align:center;padding:7px;color:#f59e0b;font-weight:600">{"⭐"*n} {n}-star rating received!</div>'

def fb_text(kind):
    m = {"helpful": ("🎉 Glad it was helpful!","#10b981"), "improve": ("💪 Will keep improving!","#f59e0b")}
    txt, col = m.get(kind, ("✅ Thanks!","#10b981"))
    return f'<div style="text-align:center;padding:7px;color:{col};font-weight:500">{txt}</div>'


# ════════════════════════════════════════════════════════════
#  CSS — UNCHANGED
# ════════════════════════════════════════════════════════════
_CSS_STATIC = """
@import url('https://fonts.googleapis.com/css2?family=Poppins:wght@400;600;700;800;900&family=Inter:wght@300;400;500;600;700&family=Space+Grotesk:wght@400;500;600;700&family=Noto+Sans:wght@400;700&family=Noto+Sans+Telugu:wght@400;700&family=Noto+Sans+Devanagari:wght@400;700&display=swap');

/* ── DESIGN TOKENS ── */
:root {
  --bg:     #06152e;
  --card:   rgba(11,31,58,0.90);
  --cyan:   #00d4ff;
  --blue:   #38bdf8;
  --teal:   #14b8a6;
  --violet: #8b5cf6;
  --text:   #f0f9ff;
  --muted:  #94a3b8;
  --border: rgba(0,212,255,0.14);
  --bhi:    rgba(0,212,255,0.42);
  --inpbg:  rgba(4,9,22,0.75);
}

/* ── LIGHT MODE ── */
body.is-light {
  --bg:     #eef4ff;
  --card:   rgba(255,255,255,0.95);
  --text:   #0f172a;
  --muted:  #64748b;
  --cyan:   #0891b2;
  --border: rgba(8,145,178,0.18);
  --bhi:    rgba(8,145,178,0.45);
  --inpbg:  rgba(248,250,255,0.9);
}

/* ── BASE ── */
*, *::before, *::after { box-sizing: border-box; }

body {
  background: var(--bg) !important;
  background-image:
    radial-gradient(ellipse 55% 38% at 12% 4%,  rgba(0,212,255,.09),  transparent),
    radial-gradient(ellipse 48% 32% at 88% 8%,  rgba(99,102,241,.07), transparent),
    radial-gradient(ellipse 38% 28% at 50% 96%, rgba(16,185,129,.05), transparent) !important;
  background-attachment: fixed !important;
  font-family: 'Inter','Noto Sans',sans-serif !important;
  color: var(--text) !important;
  min-height: 100vh !important;
  overflow-x: hidden !important;
}

/* ── STAR LAYERS (box-shadow injected separately below) ── */
.is-sl1, .is-sl2, .is-sl3 {
  position: fixed !important;
  top: 0 !important; left: 0 !important;
  width: 2px !important; height: 2px !important;
  border-radius: 50% !important;
  background: white !important;
  pointer-events: none !important;
  z-index: -10 !important;
}
@keyframes twA { 0%{opacity:.7} 100%{opacity:.2} }
@keyframes twB { 0%{opacity:.2} 100%{opacity:.6} }
@keyframes twC { 0%{opacity:.3} 100%{opacity:.9} }
.is-sl1 { animation: twA  7s ease-in-out infinite alternate; }
.is-sl2 { animation: twB 11s ease-in-out infinite alternate; }
.is-sl3 { animation: twC  5s ease-in-out infinite alternate; }
body.is-light .is-sl1,
body.is-light .is-sl2,
body.is-light .is-sl3 { opacity: .06 !important; }

/* ── GRADIO CONTAINER ── */
.gradio-container {
  max-width: 1440px !important;
  width: 97% !important;
  margin: 0 auto !important;
  padding: 0 14px 48px !important;
  background: transparent !important;
  position: relative !important;
  z-index: 1 !important;
}
.gradio-container .form,
.gradio-container .block {
  background: transparent !important;
  border: none !important;
  box-shadow: none !important;
}

/* ── ANIMATIONS ── */
@keyframes titleShimmer { 0%,100%{background-position:0% 50%} 50%{background-position:100% 50%} }
@keyframes fadeUp { from{opacity:0;transform:translateY(12px)} to{opacity:1;transform:translateY(0)} }
@keyframes pulse  { 0%,100%{box-shadow:0 0 12px rgba(0,212,255,.08)} 50%{box-shadow:0 0 24px rgba(0,212,255,.20)} }

/* ── FIXED BUTTONS ── */
#is-mute-btn, #is-theme-btn {
  position: fixed !important; top: 12px !important; z-index: 9999 !important;
  background: rgba(6,21,46,.94) !important; backdrop-filter: blur(12px);
  border: 1px solid var(--border) !important; color: var(--muted) !important;
  border-radius: 100px !important; padding: 6px 13px !important;
  font-size: .73rem !important; font-weight: 500 !important;
  cursor: pointer !important; transition: all .22s !important;
  font-family: 'Inter',sans-serif !important;
}
#is-mute-btn  { right: 132px !important; }
#is-theme-btn { right: 14px !important; }
#is-mute-btn:hover, #is-theme-btn:hover {
  border-color: var(--bhi) !important; color: var(--cyan) !important;
}

#welcome-audio {
  height: 0 !important;
  min-height: 0 !important;
  overflow: hidden !important;
  opacity: 0 !important;
  margin: 0 !important;
  padding: 0 !important;
}

/* ── HERO ── */
.is-hero { text-align:center; padding:38px 0 12px; animation:fadeUp .7s ease; }
.is-logo {
  font-family: 'Poppins',sans-serif !important;
  font-size: 3.8rem; font-weight: 950; letter-spacing: 1.5px;
  background: linear-gradient(135deg,#e0f2fe 0%,#38bdf8 22%,#00d4ff 48%,#22d3ee 68%,#e0f2fe 100%);
  background-size: 300% 300%; animation: titleShimmer 5s ease infinite;
  -webkit-background-clip: text; -webkit-text-fill-color: transparent; background-clip: text;
  display: inline-block; margin-bottom: 7px;
}
.is-sub {
  font-family: 'Space Grotesk',sans-serif !important;
  color: var(--muted); font-size: .88rem; letter-spacing: .5px; margin: 0 0 16px;
}

/* ── WELCOME HERO CARD ── */
.is-hero-card {
  max-width: 520px; margin: 0 auto 20px;
  background: rgba(11,31,58,.83); border: 1px solid rgba(0,212,255,.22);
  border-radius: 18px; padding: 20px 22px;
  backdrop-filter: blur(11px); animation: pulse 6s ease infinite;
  text-align: center;
}
body.is-light .is-hero-card { background: rgba(255,255,255,.92); border-color: rgba(8,145,178,.28); }
.is-hero-card-title { font-family:'Poppins',sans-serif !important; font-size:.92rem; font-weight:700; color:#e0f2fe; margin-bottom:4px; }
body.is-light .is-hero-card-title { color:#0f172a; }
.is-hero-tagline { color:var(--muted); font-size:.80rem; margin-bottom:11px; }
.is-hero-features { display:flex; flex-wrap:wrap; gap:6px; justify-content:center; margin-bottom:14px; }
.is-hero-feat {
  background:rgba(0,212,255,.07); border:1px solid rgba(0,212,255,.15);
  color:#67e8f9; border-radius:5px; padding:3px 9px; font-size:.69rem; font-weight:500;
}
body.is-light .is-hero-feat { color:#0891b2; background:rgba(8,145,178,.08); }
.is-get-started {
  display:inline-block; background:linear-gradient(135deg,#0891b2,#00d4ff) !important;
  color:#fff !important; border:none !important; border-radius:9px !important;
  padding:9px 22px !important; font-size:.84rem !important; font-weight:700 !important;
  box-shadow:0 4px 16px rgba(0,212,255,.35) !important; transition:all .2s !important;
  font-family:'Inter',sans-serif !important; cursor:pointer !important;
}
.is-get-started:hover { box-shadow:0 6px 22px rgba(0,212,255,.55) !important; transform:translateY(-1px) !important; }

/* ══════════════════════════════════════════════════════
   SIDEBAR — width 290px, compact, dense
   ══════════════════════════════════════════════════════ */
.is-sidebar {
  background: var(--card) !important;
  border: 1px solid var(--border) !important;
  border-radius: 12px !important;
  padding: 4px 5px !important;
  backdrop-filter: blur(14px) !important;
  transition: background .3s !important;
  width: 290px !important;
  min-width: 290px !important;
}

/* Bold compact section labels */
.is-slabel {
  font-size: .85rem !important; font-weight: 800 !important;
  letter-spacing: 1.15px !important; text-transform: uppercase !important;
  color: #00d4ff !important;
  margin-bottom: 1px !important; margin-top: 1.5px !important;
  display: block !important; line-height: 1 !important;
}
.is-slabel:first-child { margin-top: 0 !important; }
.is-sdiv { height: 1px; background: var(--border); margin: 0px 0; }

/* COMPACT horizontal radio buttons inside sidebar */
.is-sidebar .gradio-radio > div,
.is-sidebar fieldset > div,
.is-sidebar [data-testid="radio-group"] {
  display: flex !important;
  flex-direction: row !important;
  flex-wrap: wrap !important;
  gap: 1.5px !important;
  padding: 0 !important; margin: 0 !important;
}
.is-sidebar input[type="radio"] + span,
.is-sidebar .gradio-radio label {
  background: rgba(15,23,42,.95) !important;
  border: 1px solid rgba(255,255,255,.12) !important;
  border-radius: 6px !important; 
  min-height: 40px !important;
  font-size: 0.82rem !important;
  font-weight: 700 !important;
  color: #cbd5e1 !important;
  display: flex !important;
  align-items: center !important;
  padding: 8px 10px !important; 
  transition: all .25s ease !important;
  white-space: nowrap !important;
  line-height: 1.2 !important; 
  cursor: pointer !important;
}

.is-sidebar input[type="radio"]:checked + span {
  background: linear-gradient(
    135deg,
    rgba(0,212,255,.28),
    rgba(56,189,248,.22)
  ) !important;
  border: 1.5px solid #00d4ff !important;
  color: #ffffff !important;
  box-shadow:
    0 0 12px rgba(0,212,255,.45),
    0 0 24px rgba(0,212,255,.15) !important;
}


body.is-light .is-sidebar input[type="radio"] + span {
  background: rgba(240,246,255,.8) !important; border-color: rgba(8,145,178,.15) !important;
}

.lang-grid {
  width: 100% !important;
}

/* THIS is the real Gradio radio container */
.lang-grid .gr-radio {
  display: grid !important;
  grid-template-columns: repeat(3, 1fr) !important;
  gap: 3px !important;
}

/* each option */
.lang-grid label {
  display: flex !important;
  align-items: center !important;
  justify-content: center !important;
  background: rgba(4,9,22,.65) !important;
  border: 1px solid rgba(0,212,255,.2) !important;
  padding: 2px 3px !important;
  border-radius: 1.5px !important;
  cursor: pointer !important;
  font-size: 0.55rem !important;
  transition: 0.2s !important;
}

/* hover */
.lang-grid label:hover {
  border-color: rgba(0,212,255,.6) !important;
  transform: scale(1.03);
}

.is-dropdown select {
  background-color: rgba(4,9,22,.82) !important;
  color: white !important;
  border: 1px solid rgba(0,212,255,.3) !important;
  border-radius: 4px !important;
}

.gradio-accordion {
  overflow: visible !important;
}

.gradio-accordion > div {
  overflow: visible !important;
}

/* Hide "Radio" default label in sidebar */
.is-sidebar .gradio-radio .label-wrap,
.is-sidebar fieldset legend { display: none !important; }


/* ══════════════════════════════════════════════════════
   EQUAL SQUARE INPUT CARDS — height 480px, padding 18px
   ══════════════════════════════════════════════════════ */
.is-card {
  background: var(--card) !important;
  border: 2px solid var(--border) !important;
  border-radius: 15px !important;
  padding: 18px !important;
  backdrop-filter: blur(16px) !important;
  transition: border-color .22s, transform .18s !important;
  height: 480px !important;
  min-height: 480px !important;
  max-height: 480px !important;
  overflow: hidden !important;
  display: flex !important;
  flex-direction: column !important;
}
.is-card:hover { border-color: var(--bhi) !important; transform: translateY(-1px) !important; }
body.is-light .is-card { background: rgba(255,255,255,.96) !important; }

/* Coloured top borders per card */
.is-card-cyan   { border-top: 2px solid rgba(0,212,255,.65)   !important; }
.is-card-blue   { border-top: 2px solid rgba(56,189,248,.65)  !important; }
.is-card-teal   { border-top: 2px solid rgba(20,184,166,.65)  !important; }
.is-card-violet { border-top: 2px solid rgba(139,92,246,.65)  !important; }

/* Card label */
.is-clabel {
  font-size: .70rem; font-weight: 700; letter-spacing: 1.5px;
  text-transform: uppercase; display: block;
  margin-bottom: 5px; flex-shrink: 0;
}
.cl-cyan   { color:var(--cyan); }
.cl-blue   { color:var(--blue); }
.cl-teal   { color:var(--teal); }
.cl-violet { color:var(--violet); }

/* Compact textarea */
.is-card textarea {
  background: var(--inpbg) !important;
  border: 1px solid rgba(255,255,255,.07) !important;
  border-radius: 8px !important; color: var(--text) !important;
  font-size: .75rem !important; line-height: 1.5 !important;
  padding: 5px 7px !important; resize: none !important;
  font-family: 'Inter','Noto Sans','Noto Sans Telugu','Noto Sans Devanagari',sans-serif !important;
  flex: 1 !important; min-height: 200 !important; max-height: 300px !important;
}
.is-card textarea:focus {
  border-color: rgba(0,212,255,.48) !important;
  box-shadow: 0 0 0 2px rgba(0,212,255,.07) !important;
}

/* Compact file upload */
.is-card .upload-container,
.is-card [data-testid="file"] {
  background: rgba(4,9,22,.5) !important;
  border: 1.5px dashed rgba(0,212,255,.22) !important;
  border-radius: 8px !important;
  flex: 1 !important;
  min-height: 0 !important; max-height: 62px !important;
  display: flex !important; align-items: center !important;
  justify-content: center !important;
  overflow: hidden !important;
}
.is-card .upload-container:hover { border-color: rgba(0,212,255,.45) !important; }
.is-card .upload-container span,
.is-card .upload-container p { font-size: .68rem !important; color: var(--muted) !important; }

/* Compact audio widget */
.is-card .waveform-container,
.is-card [data-testid="audio"],
.is-card audio {
  flex: 1 !important; min-height: 0 !important;
  max-height: 60px !important; border-radius: 8px !important;
}
.is-card audio { width: 100% !important; height: 32px !important; }

/* Card button row — always at bottom */
.is-card-btns {
  display: flex !important; gap: 5px !important;
  margin-top: 5px !important; flex-shrink: 0 !important;
}

/* ── ROW GAP between the two card rows ── */
#is-card-grid { gap: 12px !important; }
#is-row1, #is-row2 { gap: 12px !important; margin-bottom: 0 !important; }
/* Kill any extra padding Gradio adds to rows */
#is-row1 > div, #is-row2 > div { padding: 0 !important; }

/* ── Global label cleanup ── */
.gradio-container label,
.gradio-container .label-wrap span {
  color: var(--muted) !important; font-size: .77rem !important; font-weight: 500 !important;
}

/* ── SUMMARY OUTPUT ── */
#is-sum textarea {
  min-height: 200px !important; max-height: 440px !important;
  background: var(--inpbg) !important;
  border: 1px solid rgba(0,212,255,.2) !important; border-radius: 13px !important;
  color: var(--text) !important; font-size: .90rem !important; line-height: 1.8 !important;
  padding: 16px !important;
  font-family: 'Inter','Noto Sans','Noto Sans Telugu','Noto Sans Devanagari',sans-serif !important;
}
#is-notes textarea {
  min-height: 125px !important; background: var(--inpbg) !important;
  border: 1px solid rgba(20,184,166,.22) !important; border-radius: 11px !important;
  color: #a7f3d0 !important; font-size: .85rem !important; line-height: 1.68 !important;
  padding: 11px !important; font-family: 'Inter',monospace !important;
}
body.is-light #is-notes textarea { color: #0d9488 !important; }
#is-trans textarea {
  min-height: 145px !important; background: var(--inpbg) !important;
  border: 1px solid rgba(139,92,246,.22) !important; border-radius: 11px !important;
  color: var(--text) !important; font-size: .88rem !important; line-height: 1.8 !important;
  padding: 14px !important;
  font-family: 'Inter','Noto Sans','Noto Sans Telugu','Noto Sans Devanagari',sans-serif !important;
}
.gr-textbox textarea::-webkit-scrollbar {
    width: 8px;
}

.gr-textbox textarea::-webkit-scrollbar-thumb {
    background: #00D4FF;
    border-radius: 10px;
}

.gr-textbox textarea::-webkit-scrollbar-track {
    background: rgba(255,255,255,0.08);
}

/* ── STAT BOXES ── */
.is-stat {
  background: rgba(7,20,38,.8) !important;
  border: 1px solid rgba(255,255,255,.06) !important;
  border-radius: 11px !important; padding: 14px 8px !important;
  text-align: center !important; transition: all .2s !important;
}
.is-stat:hover { border-color: rgba(0,212,255,.28) !important; transform: translateY(-2px) !important; }
.is-stat input, .is-stat textarea {
  background: transparent !important; border: none !important; text-align: center !important;
  font-size: 1.3rem !important; font-weight: 700 !important; color: #67e8f9 !important;
  padding: 0 !important; box-shadow: none !important; font-family: 'Poppins',sans-serif !important;
}
.is-stat .label-wrap span { color: var(--muted) !important; font-size: .68rem !important; }
body.is-light .is-stat { background: rgba(240,246,255,.9) !important; }
body.is-light .is-stat input,
body.is-light .is-stat textarea { color: #0891b2 !important; }

/* ── COMPARE ── */
.is-cmp textarea {
  min-height: 255px !important; background: var(--inpbg) !important;
  border: 1px solid rgba(255,255,255,.07) !important; border-radius: 11px !important;
  color: var(--text) !important; font-size: .85rem !important; line-height: 1.7 !important;
  padding: 13px !important;
  font-family: 'Inter','Noto Sans','Noto Sans Telugu',sans-serif !important;
}

/* ── BUTTONS ── */
.is-btn {
  background: linear-gradient(135deg,#0891b2,#00d4ff) !important;
  color: #fff !important; border: none !important; border-radius: 8px !important;
  font-weight: 600 !important; font-size: .80rem !important; padding: 7px 16px !important;
  box-shadow: 0 3px 12px rgba(0,212,255,.3) !important; transition: all .2s !important;
  font-family: 'Inter',sans-serif !important;
}
.is-btn:hover { box-shadow: 0 5px 18px rgba(0,212,255,.5) !important; transform: translateY(-1px) !important; }
.is-btn-teal   { background: linear-gradient(135deg,#0d9488,#14b8a6) !important; box-shadow: 0 3px 12px rgba(20,184,166,.25) !important; }
.is-btn-violet { background: linear-gradient(135deg,#6d28d9,#8b5cf6) !important; box-shadow: 0 3px 12px rgba(139,92,246,.25) !important; }
.is-btn-blue   { background: linear-gradient(135deg,#0369a1,#38bdf8) !important; box-shadow: 0 3px 12px rgba(56,189,248,.25) !important; }
.is-btn-green  { background: linear-gradient(135deg,#059669,#10b981) !important; box-shadow: 0 3px 12px rgba(16,185,129,.25) !important; }
.is-btn-ghost  {
  background: rgba(255,255,255,.04) !important; color: var(--muted) !important;
  border: 1px solid rgba(255,255,255,.09) !important; border-radius: 7px !important;
  font-weight: 500 !important; box-shadow: none !important; padding: 7px 12px !important;
}
.is-btn-ghost:hover { background: rgba(255,255,255,.09) !important; color: var(--text) !important; }

/* ── ACCORDION ── */
.is-acc {
  background: var(--card) !important; border: 1px solid var(--border) !important;
  border-radius: 12px !important; margin-bottom: 9px !important;
  backdrop-filter: blur(16px) !important; overflow: hidden !important;
  transition: border-color .2s, background .3s !important;
}
.is-acc:hover { border-color: var(--bhi) !important; }
.is-acc > .label-wrap {
  background: transparent !important; padding: 13px 17px !important;
  cursor: pointer !important; transition: background .14s !important;
}
.is-acc > .label-wrap:hover { background: rgba(0,212,255,.04) !important; }
.is-acc > .label-wrap span {
  color: var(--text) !important; font-size: .85rem !important;
  font-weight: 600 !important; font-family: 'Space Grotesk',sans-serif !important;
}
.is-acc > .wrap { padding: 0 17px 17px !important; }

/* ── GENERAL RADIO ── */
.gradio-container input[type="radio"] + span {
  background: rgba(5,8,22,.5) !important; border: 1px solid rgba(255,255,255,.07) !important;
  border-radius: 7px !important; color: var(--muted) !important;
  padding: 5px 10px !important; font-size: .77rem !important; transition: all .2s !important;
}
.gradio-container input[type="radio"]:checked + span {
  background: rgba(0,212,255,.12) !important; border-color: rgba(0,212,255,.45) !important;
  color: #67e8f9 !important;
}

/* ── FILE UPLOAD outside cards ── */
.gradio-container .upload-container,
.gradio-container [data-testid="file"] {
  background: rgba(4,9,22,.5) !important;
  border: 2px dashed rgba(0,212,255,.18) !important; border-radius: 10px !important;
}
.gradio-container .upload-container:hover { border-color: rgba(0,212,255,.4) !important; }
.gradio-container audio { border-radius: 9px !important; }

/* ── MD TABLE ── */
.gradio-container table { background: rgba(7,20,38,.7) !important; border-radius: 9px !important; overflow: hidden !important; width: 100% !important; border-collapse: collapse !important; }
.gradio-container th { background: rgba(0,212,255,.12) !important; color: #67e8f9 !important; padding: 8px 11px !important; font-size: .77rem !important; font-weight: 600 !important; }
.gradio-container td { color: var(--muted) !important; padding: 7px 11px !important; border-top: 1px solid rgba(255,255,255,.04) !important; font-size: .77rem !important; }

/* ── RESULTS ── */
.is-results-wrap { text-align:center; padding:20px 0 13px; }
.is-results-title {
  font-family:'Poppins',sans-serif; font-size:1.4rem; font-weight:800;
  background:linear-gradient(135deg,#e0f2fe,#38bdf8,#00d4ff);
  -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text;
  display:inline-block; margin-bottom:5px;
}
.is-results-line { width:130px; height:2px; margin:0 auto; background:linear-gradient(90deg,transparent,#00d4ff,#38bdf8,#00d4ff,transparent); border-radius:2px; }

/* ── HISTORY / FEEDBACK / FOOTER ── */
.is-hist-card { background:var(--card); border:1px solid var(--border); border-radius:13px; padding:15px; margin-bottom:11px; transition:background .3s; }
.is-hist-title { font-family:'Space Grotesk',sans-serif; font-size:.87rem; font-weight:700; color:var(--text); margin-bottom:9px; }
.is-fb-card { background:var(--card); border:1px solid var(--border); border-radius:13px; padding:17px; margin-bottom:11px; text-align:center; transition:background .3s; }
.is-fb-title { font-family:'Poppins',sans-serif; font-size:.92rem; font-weight:700; color:var(--text); margin-bottom:3px; }
.is-fb-sub { color:var(--muted); font-size:.79rem; margin-bottom:11px; }
.is-footer { text-align:center; padding:24px 0 10px; }
.is-footer-line { height:1px; width:100%; max-width:330px; margin:0 auto 15px; background:linear-gradient(90deg,transparent,rgba(0,212,255,.3),transparent); }
.is-footer-title { font-family:'Poppins',sans-serif; font-size:.93rem; font-weight:700; background:linear-gradient(135deg,#e0f2fe,#00d4ff); -webkit-background-clip:text; -webkit-text-fill-color:transparent; background-clip:text; display:inline-block; margin-bottom:3px; }
.is-footer-sub { color:var(--muted); font-size:.78rem; margin-bottom:6px; }
.is-footer-brand { color:rgba(71,85,105,.5); font-size:.70rem; letter-spacing:.5px; }
.is-footer-brand span { color:rgba(0,212,255,.4); }
.is-sec-lbl { font-size:.59rem; font-weight:700; letter-spacing:2px; text-transform:uppercase; color:var(--cyan); display:block; margin-bottom:6px; }
"""

_CSS_STARS = """
.is-sl1 {{ box-shadow: {A}; }}
.is-sl2 {{ box-shadow: {B}; }}
.is-sl3 {{ box-shadow: {C}; }}
""".format(A=STARS_A, B=STARS_B, C=STARS_C)

CSS = _CSS_STATIC + _CSS_STARS

THEME = gr.themes.Base(
    primary_hue="cyan", secondary_hue="sky", neutral_hue="slate",
    font=[gr.themes.GoogleFont("Inter"),"Noto Sans","sans-serif"],
)

# ════════════════════════════════════════════════════════════
#  BUILD UI — UNCHANGED
# ════════════════════════════════════════════════════════════
with gr.Blocks(title="InstaSummary — AI Summarizer") as demo:

    gr.HTML('<div class="is-sl1"></div><div class="is-sl2"></div><div class="is-sl3"></div>')

    gr.HTML("""
    <button id="is-mute-btn" data-muted="0" onclick="(function(btn){
      var m = btn.getAttribute('data-muted')==='1';
      document.querySelectorAll('audio').forEach(function(a){ a.muted=!m; });
      btn.setAttribute('data-muted', m?'0':'1');
      btn.textContent = m ? '🔇 Mute' : '🔊 Unmute';
    })(this)">🔇 Mute</button>

    <button id="is-theme-btn" onclick="(function(btn){
      document.body.classList.toggle('is-light');
      btn.textContent = document.body.classList.contains('is-light') ? '🌙 Dark' : '☀️ Light';
    })(this)">☀️ Light</button>""")

    welcome_audio = gr.Audio(
        visible=False, autoplay=False, interactive=False,
        show_label=False, elem_id="welcome-audio")

    demo.load(fn=_welcome_audio_fn, inputs=[], outputs=[welcome_audio])

    gr.HTML("""
    <div class="is-hero">
      <div class="is-logo">⚡ InstaSummary</div>
      <p class="is-sub">AI Powered &nbsp;•&nbsp; Context Aware &nbsp;•&nbsp; Smart Summarizer System</p>
    </div>""")

    gr.HTML("""
    <div class="is-hero-card">
      <div style="position:relative;">
        <div class="is-hero-card-title" style="text-align:center;width:100%;">
          🚀 Welcome to InstaSummary
        </div>
      </div>
      <div class="is-hero-tagline">Transform long documents into smart insights using AI.</div>
      <div class="is-hero-features">
        <span class="is-hero-feat">✓ AI Summarization</span>
        <span class="is-hero-feat">✓ 15 Languages</span>
        <span class="is-hero-feat">✓ Voice Output</span>
        <span class="is-hero-feat">✓ PDF 300+ pages</span>
        <span class="is-hero-feat">✓ Audio Transcription</span>
        <span class="is-hero-feat">✓ CSV / Excel</span>
        <span class="is-hero-feat">✓ PDF &amp; DOCX Export</span>
      </div>
      <button class="is-get-started"
        onclick="document.getElementById('is-main').scrollIntoView({behavior:'smooth'})">
        ⚡ Start Summarizing
      </button>
    </div>""")

    gr.HTML('<div id="is-main"></div>')

    with gr.Row(equal_height=True):

        with gr.Column(scale=2, min_width=280, elem_classes="is-sidebar"):
            gr.HTML('<span class="is-slabel" style="margin-top:0">🤖 AI MODEL</span>')
            model_sel = gr.Radio(
                ["BERT"],
                value="BERT", label="", show_label=False)

            gr.HTML('<span class="is-slabel">📏 SUMMARY LENGTH</span>')
            length_sel = gr.Radio(
                ["Short (50-80)","Medium (100-150)","Long (200-280)"],
                value="Medium (100-150)", label="", show_label=False)

            gr.HTML('<span class="is-slabel">🎭 TONE</span>')
            tone_sel = gr.Radio(
                ["Professional","Simple English"],
                value="Professional", label="", show_label=False)

            gr.HTML('<span class="is-slabel">📋 FORMAT</span>')
            fmt_sel = gr.Radio(
                ["Paragraph","Bullet Points"],
                value="Paragraph", label="", show_label=False)

            gr.HTML('<span class="is-slabel">🌐 OUTPUT LANGUAGE</span>')
            lang_sel = gr.Radio(
                choices=LANG_CHOICES,
                value="English",
                elem_classes=["lang-grid"])

        with gr.Column(scale=5, elem_id="is-card-grid"):

            with gr.Row(equal_height=True, elem_id="is-row1"):
                with gr.Column(scale=1, elem_classes="is-card is-card-cyan"):
                    gr.HTML('<span class="is-clabel cl-cyan">📝 Text Input</span>')
                    text_in = gr.Textbox(
                        placeholder="✏️ Paste article, paper, legal doc, code…",
                        lines=18, max_lines=60, label="", show_label=False, container=False)
                    with gr.Row(elem_classes="is-card-btns"):
                        text_btn = gr.Button("⚡ Summarize", elem_classes=["is-btn"], scale=3)
                        clr_text = gr.Button("✕", elem_classes=["is-btn-ghost"], scale=1)

                with gr.Column(scale=1, elem_classes="is-card is-card-blue"):
                    gr.HTML('<span class="is-clabel cl-blue">📄 PDF Upload</span>')
                    pdf_in = gr.File(file_types=[".pdf"], type="filepath",
                                     label="", show_label=False, container=False)
                    with gr.Row(elem_classes="is-card-btns"):
                        pdf_btn = gr.Button("📄 Summarize PDF",
                                             elem_classes=["is-btn","is-btn-blue"], scale=3)
                        clr_pdf = gr.Button("✕", elem_classes=["is-btn-ghost"], scale=1)

            with gr.Row(equal_height=True, elem_id="is-row2"):
                with gr.Column(scale=1, elem_classes="is-card is-card-teal"):
                    gr.HTML('<span class="is-clabel cl-teal">🎙️ Audio Input</span>')
                    audio_in = gr.Audio(type="filepath", sources=["upload","microphone"],
                                        label="", show_label=False, container=False)
                    with gr.Row(elem_classes="is-card-btns"):
                        audio_btn = gr.Button("🎙️ Transcribe & Summarize",
                                               elem_classes=["is-btn","is-btn-teal"], scale=3)
                        clr_audio = gr.Button("✕", elem_classes=["is-btn-ghost"], scale=1)

                with gr.Column(scale=1, elem_classes="is-card is-card-violet"):
                    gr.HTML('<span class="is-clabel cl-violet">📊 CSV / Excel</span>')
                    csv_in = gr.File(file_types=[".csv",".xlsx",".xls"],
                                     type="filepath", label="", show_label=False, container=False)
                    with gr.Row(elem_classes="is-card-btns"):
                        csv_btn = gr.Button("📊 Analyze & Summarize",
                                             elem_classes=["is-btn","is-btn-violet"], scale=3)
                        clr_csv = gr.Button("✕", elem_classes=["is-btn-ghost"], scale=1)

    status_html = gr.HTML(value=_status(
        "Ready — paste text or upload a file, then click Summarize.", "info"))

    info_html = gr.HTML(value="")
    
    gr.HTML("""
    <div class="is-results-wrap">
      <div class="is-results-title">✨ Results</div>
      <div class="is-results-line"></div>
    </div>""")

    with gr.Accordion("✨  AI Summary  ·  Voice  ·  Translation  ·  Smart Notes",
                       open=True, elem_classes="is-acc"):
        title_html  = gr.HTML(value="")
        summary_out = gr.Textbox(label="✨ AI Summary", lines=18, max_lines=40,
                                  interactive=False, elem_id="is-sum")
        gr.HTML('<div style="margin-top:12px;margin-bottom:4px;font-size:.59rem;font-weight:700;'
                'letter-spacing:2px;text-transform:uppercase;color:#00d4ff">🔊 Voice Output</div>')
        voice_out = gr.Audio(label="", interactive=False, autoplay=True)
        gr.HTML("""
        <div style="margin-top:14px;margin-bottom:4px;font-size:.59rem;font-weight:700;
          letter-spacing:2px;text-transform:uppercase;color:#8b5cf6">
          🌍 Translate Summary</div>
        <div style="color:#4b5563;font-size:.7rem;margin-bottom:7px">
          Select a language → translated text &amp; voice are generated automatically</div>""")
        with gr.Row():
            acc_trans_lang = gr.Radio(
                choices=LANG_CHOICES, value="Hindi",
                label="Translate To", scale=2, elem_classes=["lang-grid"])
            acc_trans_btn = gr.Button("🌍 Translate & Speak",
                                       elem_classes=["is-btn","is-btn-violet"], scale=1)
        trans_out   = gr.Textbox(label="Translated Output", lines=5,
                                  interactive=False, elem_id="is-trans")
        trans_voice = gr.Audio(label="Translated Voice", interactive=False)
        gr.HTML('<div style="margin-top:14px;margin-bottom:4px;font-size:.59rem;font-weight:700;'
                'letter-spacing:2px;text-transform:uppercase;color:#14b8a6">🧠 Smart Notes</div>')
        notes_out = gr.Textbox(label="", lines=5, max_lines=14,
                                interactive=False, elem_id="is-notes", container=False)
        ty_html = gr.HTML(value="")

    with gr.Accordion("📈  Statistics", open=False, elem_classes="is-acc"):
        gr.HTML('<span class="is-sec-lbl">Original Document</span>')
        with gr.Row():
            with gr.Column(elem_classes="is-stat"): wb_out=gr.Textbox(label="Words",     interactive=False,lines=1)
            with gr.Column(elem_classes="is-stat"): sb_out=gr.Textbox(label="Sentences", interactive=False,lines=1)
            with gr.Column(elem_classes="is-stat"): rb_out=gr.Textbox(label="Read Time", interactive=False,lines=1)
        gr.HTML('<span class="is-sec-lbl" style="margin-top:10px">AI Summary</span>')
        with gr.Row():
            with gr.Column(elem_classes="is-stat"): wa_out=gr.Textbox(label="Words",     interactive=False,lines=1)
            with gr.Column(elem_classes="is-stat"): sa_out=gr.Textbox(label="Sentences", interactive=False,lines=1)
            with gr.Column(elem_classes="is-stat"): ra_out=gr.Textbox(label="Read Time", interactive=False,lines=1)
        with gr.Row():
            with gr.Column(elem_classes="is-stat"):
                ratio_out=gr.Textbox(label="📉 Compression %",interactive=False,lines=1)

    with gr.Accordion("🔍  Compare — Original vs Summary", open=False, elem_classes="is-acc"):
        with gr.Row(equal_height=True):
            with gr.Column(scale=1):
                gr.HTML('<div style="color:#94a3b8;font-size:.72rem;font-weight:600;margin-bottom:5px">📄 Original Text</div>')
                orig_cmp=gr.Textbox(lines=10,interactive=False,show_label=False,
                    elem_classes=["is-cmp"],placeholder="Original text appears here after summarizing…")
            with gr.Column(scale=1):
                gr.HTML('<div style="color:#67e8f9;font-size:.72rem;font-weight:600;margin-bottom:5px">✨ AI Summary</div>')
                summ_cmp=gr.Textbox(lines=10,interactive=False,show_label=False,
                    elem_classes=["is-cmp"],placeholder="Summary appears here after summarizing…")

    with gr.Accordion("🔑  Keywords", open=False, elem_classes="is-acc"):
        kw_html=gr.HTML(value='<p style="color:#374151;padding:7px;font-size:.83rem;margin:0">Keywords appear after summarizing.</p>')

    with gr.Accordion("📊  Data Analysis — CSV / Excel Only", open=False, elem_classes="is-acc"):
        gr.HTML('<div style="background:rgba(0,212,255,.04);border:1px solid rgba(0,212,255,.1);'
                'border-radius:8px;padding:8px 12px;margin-bottom:10px;color:#94a3b8;font-size:.78rem">'
                '<strong style="color:#22d3ee">ℹ️</strong> Only available for CSV and Excel files.</div>')
        data_preview  = gr.Markdown(value="*Upload a CSV or Excel file and click Analyze.*")
        gr.HTML('<span class="is-sec-lbl" style="margin-top:10px">💡 Auto Insights</span>')
        data_insights = gr.Textbox(lines=4,interactive=False,show_label=False,
                                    placeholder="Trend insights appear here…")

    with gr.Accordion("📥  Download Reports — PDF + DOCX", open=False, elem_classes="is-acc"):
        gr.HTML('<div style="background:rgba(0,212,255,.04);border:1px solid rgba(0,212,255,.1);'
                'border-radius:8px;padding:9px 13px;margin-bottom:11px;color:#94a3b8;font-size:.78rem;line-height:1.7">'
                '<strong style="color:#22d3ee">📄 PDF</strong> — Original, AI summary, smart notes, keywords, stats, translation<br>'
                '<strong style="color:#22d3ee">📝 DOCX</strong> — Same content, editable in Microsoft Word</div>')
        with gr.Row():
            pdf_out  = gr.File(label="📄 PDF Report",  interactive=False)
            docx_out = gr.File(label="📝 DOCX Report", interactive=False)

    gr.HTML('<div class="is-hist-card"><div class="is-hist-title">📚 Session History</div>')
    history_html  = gr.HTML(value=_hist_html(_load_hist()))
    history_state = gr.State(value=_load_hist())
    gr.HTML('</div>')

    gr.HTML('<div class="is-fb-card"><div class="is-fb-title">⭐ Rate Your Experience</div>'
            '<div class="is-fb-sub">How was your experience?</div></div>')
    with gr.Row():
        fb1=gr.Button("⭐1",elem_classes=["is-btn-ghost"],scale=1)
        fb2=gr.Button("⭐2",elem_classes=["is-btn-ghost"],scale=1)
        fb3=gr.Button("⭐3",elem_classes=["is-btn-ghost"],scale=1)
        fb4=gr.Button("⭐4",elem_classes=["is-btn-ghost"],scale=1)
        fb5=gr.Button("⭐5",elem_classes=["is-btn-ghost"],scale=1)
    with gr.Row():
        fb_helpful=gr.Button("👍 Helpful",           elem_classes=["is-btn-green"], scale=1)
        fb_improve=gr.Button("👎 Needs Improvement", elem_classes=["is-btn-red"],   scale=1)
    fb_html=gr.HTML(value="")

    gr.HTML("""
    <div class="is-footer">
      <div class="is-footer-line"></div>
      <div class="is-footer-title">❤️ Thank You For Using InstaSummary ❤️</div>
      <div class="is-footer-sub">Transforming Long Information Into Smart Knowledge</div>
      <div class="is-footer-brand">
        Made with AI &nbsp;·&nbsp; <span>Gradio</span> &nbsp;·&nbsp;
        <span>Hugging Face</span> &nbsp;·&nbsp; <span>Transformers</span> &nbsp;·&nbsp;
        <span>Python</span>
      </div>
    </div>""")

    ALL_OUT=[
        status_html,   # 0
        info_html,     # 1
        title_html,    # 2
        summary_out,   # 3
        voice_out,     # 4
        wb_out,        # 5
        sb_out,        # 6
        rb_out,        # 7
        wa_out,        # 8
        sa_out,        # 9
        ra_out,        # 10
        ratio_out,     # 11
        kw_html,       # 12
        orig_cmp,      # 13
        summ_cmp,      # 14
        data_preview,  # 15
        data_insights, # 16
        pdf_out,       # 17
        docx_out,      # 18
        ty_html,       # 19
        notes_out,     # 20
    ]
    SHARED=[model_sel,length_sel,tone_sel,lang_sel,fmt_sel]

    def _after(s, d, h): return update_history(s, d, h)

    text_btn.click( fn=process_text,  inputs=[text_in]+SHARED,  outputs=ALL_OUT).then(fn=_after,inputs=[summary_out,title_html,history_state],outputs=[history_state,history_html])
    pdf_btn.click(  fn=process_pdf,   inputs=[pdf_in]+SHARED,   outputs=ALL_OUT).then(fn=_after,inputs=[summary_out,title_html,history_state],outputs=[history_state,history_html])
    audio_btn.click(fn=process_audio, inputs=[audio_in]+SHARED, outputs=ALL_OUT).then(fn=_after,inputs=[summary_out,title_html,history_state],outputs=[history_state,history_html])
    csv_btn.click(  fn=process_csv,   inputs=[csv_in]+SHARED,   outputs=ALL_OUT).then(fn=_after,inputs=[summary_out,title_html,history_state],outputs=[history_state,history_html])

    acc_trans_btn.click(fn=translate_and_voice,
                        inputs=[summary_out, acc_trans_lang],
                        outputs=[trans_out, trans_voice])

    fb1.click(fn=lambda: fb_star(1),outputs=[fb_html])
    fb2.click(fn=lambda: fb_star(2),outputs=[fb_html])
    fb3.click(fn=lambda: fb_star(3),outputs=[fb_html])
    fb4.click(fn=lambda: fb_star(4),outputs=[fb_html])
    fb5.click(fn=lambda: fb_star(5),outputs=[fb_html])
    fb_helpful.click(fn=lambda: fb_text("helpful"),outputs=[fb_html])
    fb_improve.click(fn=lambda: fb_text("improve"),outputs=[fb_html])

    clr_text.click( fn=lambda:"",   outputs=[text_in])
    clr_pdf.click(  fn=lambda:None, outputs=[pdf_in])
    clr_audio.click(fn=lambda:None, outputs=[audio_in])
    clr_csv.click(  fn=lambda:None, outputs=[csv_in])

    demo.load(fn=_welcome_audio_fn, inputs=None, outputs=[welcome_audio])


if __name__ == "__main__":
    demo.launch(
        server_name="0.0.0.0",
        server_port=7860,
        share=False,
        ssr_mode=False,
        theme=THEME,
        css=CSS,
    )